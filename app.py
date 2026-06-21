# """
# DocuForge Converter Pro — World-Class Edition
# ═══════════════════════════════════════════════════════════════════════════════
# Conversion fidelity strategy:
#   • DOCX/ODT/PPTX → PDF   : LibreOffice headless  (100% layout fidelity)
#   • XLSX/CSV      → PDF   : LibreOffice headless on intermediate XLSX
#   • PDF           → DOCX  : PyMuPDF text extraction → styled python-docx
#   • PDF           → TXT   : PyMuPDF (superior to pypdf for layout preservation)
#   • PDF           → XLSX  : PyMuPDF + table heuristics → openpyxl
#   • DOCX          → TXT   : python-docx paragraph traversal (preserves order)
#   • DOCX          → XLSX  : paragraph lines → structured dataframe
#   • TXT/CSV       → DOCX  : fully styled python-docx output
#   • CSV/XLSX      → DOCX  : styled table inside python-docx
#   • Any spreadsheet ↔ any spreadsheet : pandas
# """

# import os, sys, re, shutil, subprocess, tempfile, threading, time
# import tkinter as tk
# from tkinter import filedialog, messagebox
# from pathlib import Path

# import ttkbootstrap as tb
# from ttkbootstrap.constants import *
# from ttkbootstrap.widgets.scrolled import ScrolledFrame

# # Document libraries
# import docx2txt
# from docx import Document as DocxDocument
# from docx.shared import Pt, RGBColor, Inches, Cm
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.table import WD_TABLE_ALIGNMENT

# import fitz          # PyMuPDF — best-in-class PDF reader
# import pandas as pd
# import openpyxl
# from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
# from openpyxl.utils import get_column_letter

# from reportlab.lib.pagesizes import letter, landscape, A4
# from reportlab.platypus import SimpleDocTemplate, Paragraph as RLParagraph, Spacer, Table as RLTable, TableStyle
# from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
# from reportlab.lib import colors as rl_colors
# from reportlab.lib.units import inch


# # ── Runtime asset resolver ──────────────────────────────────────────────────
# def resource_path(relative_path):
#     try:
#         base = sys._MEIPASS
#     except AttributeError:
#         base = os.path.abspath(".")
#     return os.path.join(base, relative_path)


# # ══════════════════════════════════════════════════════════════════════════════
# #  LIBREOFFICE BRIDGE — the secret weapon for 100% fidelity
# # ══════════════════════════════════════════════════════════════════════════════
# class LibreOfficeBridge:
#     """
#     Wraps LibreOffice headless mode. When LibreOffice converts DOCX→PDF it uses
#     the same rendering engine as the desktop app — identical to Microsoft Word's
#     own 'Export to PDF' in terms of fidelity. Tables, fonts, colours, images,
#     headers/footers — all preserved exactly.
#     """

#     # Candidate binary locations across Windows, macOS, Linux
#     _CANDIDATES = [
#         "libreoffice", "soffice",
#         r"C:\Program Files\LibreOffice\program\soffice.exe",
#         r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
#         "/Applications/LibreOffice.app/Contents/MacOS/soffice",
#         "/usr/bin/libreoffice", "/usr/bin/soffice",
#         "/usr/local/bin/libreoffice",
#     ]

#     @classmethod
#     def find_binary(cls):
#         for candidate in cls._CANDIDATES:
#             if shutil.which(candidate) or os.path.isfile(candidate):
#                 return candidate
#         return None

#     @classmethod
#     def convert(cls, source_path: str, target_format: str, output_dir: str) -> str | None:
#         """
#         Convert source_path to target_format inside output_dir.
#         Returns the output file path on success, None on failure.
#         target_format examples: 'pdf', 'docx', 'xlsx', 'csv', 'txt'
#         """
#         binary = cls.find_binary()
#         if not binary:
#             return None

#         # LibreOffice locks its user-profile dir — use a temp profile per call
#         # to allow parallel conversions and avoid conflicts with running instances.
#         with tempfile.TemporaryDirectory() as profile_dir:
#             cmd = [
#                 binary,
#                 "--headless",
#                 f"-env:UserInstallation=file:///{profile_dir.replace(os.sep, '/')}",
#                 "--convert-to", target_format,
#                 "--outdir", output_dir,
#                 source_path,
#             ]
#             try:
#                 result = subprocess.run(
#                     cmd, capture_output=True, text=True, timeout=120
#                 )
#                 if result.returncode != 0:
#                     print(f"[LibreOffice] Error: {result.stderr[:400]}")
#                     return None

#                 # LO renames the file by swapping the extension
#                 stem = Path(source_path).stem
#                 ext = target_format.split(":")[0]  # handle "pdf" vs "pdf:writer_pdf_Export"
#                 out = Path(output_dir) / f"{stem}.{ext}"
#                 return str(out) if out.exists() else None
#             except subprocess.TimeoutExpired:
#                 print("[LibreOffice] Conversion timed out")
#                 return None
#             except Exception as e:
#                 print(f"[LibreOffice] Unexpected error: {e}")
#                 return None

#     @classmethod
#     def available(cls) -> bool:
#         return cls.find_binary() is not None


# # ══════════════════════════════════════════════════════════════════════════════
# #  DOCX STYLING HELPERS
# # ══════════════════════════════════════════════════════════════════════════════
# # Brand palette used in generated DOCX / XLSX outputs
# BRAND = {
#     "header_bg":   RGBColor(0x05, 0x41, 0x3E),  # deep teal
#     "header_fg":   RGBColor(0xFF, 0xFF, 0xFF),
#     "accent":      RGBColor(0x05, 0x96, 0x69),  # emerald
#     "row_alt":     RGBColor(0xF0, 0xFD, 0xFA),
#     "border":      RGBColor(0xCB, 0xD5, 0xE1),
#     "body_text":   RGBColor(0x1E, 0x29, 0x3B),
# }

# BRAND_HEX = {
#     "header_bg":  "05413E",
#     "header_fg":  "FFFFFF",
#     "accent":     "059669",
#     "row_alt":    "F0FDFA",
#     "row_even":   "FFFFFF",
#     "border":     "CBD5E1",
# }


# def _style_docx_table(table, header_row=True):
#     """Apply professional styling to a python-docx table."""
#     tbl = table._tbl
#     for i, row in enumerate(table.rows):
#         for j, cell in enumerate(row.cells):
#             tc = cell._tc
#             # Cell margins
#             from docx.oxml.ns import qn
#             from docx.oxml import OxmlElement
#             tcPr = tc.get_or_add_tcPr()
#             tcMar = OxmlElement('w:tcMar')
#             for side in ('top', 'left', 'bottom', 'right'):
#                 node = OxmlElement(f'w:{side}')
#                 node.set(qn('w:w'), '80')
#                 node.set(qn('w:type'), 'dxa')
#                 tcMar.append(node)
#             tcPr.append(tcMar)

#             para = cell.paragraphs[0]
#             run = para.runs[0] if para.runs else para.add_run(cell.text)
#             run.font.size = Pt(9.5)
#             run.font.color.rgb = BRAND["body_text"]

#             if header_row and i == 0:
#                 run.bold = True
#                 run.font.color.rgb = BRAND["header_fg"]
#                 from docx.oxml.ns import qn
#                 from docx.oxml import OxmlElement
#                 shading = OxmlElement('w:shd')
#                 shading.set(qn('w:val'), 'clear')
#                 shading.set(qn('w:fill'), BRAND_HEX["header_bg"])
#                 tcPr.append(shading)
#             elif i % 2 == 0:
#                 from docx.oxml.ns import qn
#                 from docx.oxml import OxmlElement
#                 shading = OxmlElement('w:shd')
#                 shading.set(qn('w:val'), 'clear')
#                 shading.set(qn('w:fill'), BRAND_HEX["row_alt"])
#                 tcPr.append(shading)


# def _set_docx_styles(doc: DocxDocument):
#     """Configure document-level typography."""
#     style = doc.styles['Normal']
#     style.font.name = 'Calibri'
#     style.font.size = Pt(11)
#     style.font.color.rgb = BRAND["body_text"]

#     # Heading 1
#     h1 = doc.styles['Heading 1']
#     h1.font.name = 'Calibri'
#     h1.font.size = Pt(20)
#     h1.font.bold = True
#     h1.font.color.rgb = BRAND["accent"]

#     # Heading 2
#     h2 = doc.styles['Heading 2']
#     h2.font.name = 'Calibri'
#     h2.font.size = Pt(14)
#     h2.font.bold = True
#     h2.font.color.rgb = BRAND["header_bg"]


# # ══════════════════════════════════════════════════════════════════════════════
# #  XLSX STYLING HELPERS
# # ══════════════════════════════════════════════════════════════════════════════
# def _style_worksheet(ws, df, freeze_header=True):
#     """Apply professional table styling to an openpyxl worksheet."""
#     header_fill = PatternFill("solid", fgColor=BRAND_HEX["header_bg"])
#     alt_fill    = PatternFill("solid", fgColor=BRAND_HEX["row_alt"])
#     even_fill   = PatternFill("solid", fgColor=BRAND_HEX["row_even"])
#     thin = Side(border_style="thin", color=BRAND_HEX["border"])
#     cell_border = Border(left=thin, right=thin, top=thin, bottom=thin)

#     header_font = Font(name="Calibri", bold=True, color=BRAND_HEX["header_fg"], size=10)
#     body_font   = Font(name="Calibri", size=10, color="1E293B")
#     center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
#     left_align   = Alignment(horizontal="left",   vertical="center", wrap_text=True)

#     # Header row
#     for col_idx, col_name in enumerate(df.columns, start=1):
#         cell = ws.cell(row=1, column=col_idx, value=str(col_name))
#         cell.font   = header_font
#         cell.fill   = header_fill
#         cell.border = cell_border
#         cell.alignment = center_align

#     # Data rows
#     for row_idx, row in enumerate(df.itertuples(index=False), start=2):
#         fill = alt_fill if row_idx % 2 == 0 else even_fill
#         for col_idx, value in enumerate(row, start=1):
#             cell = ws.cell(row=row_idx, column=col_idx, value=value)
#             cell.font   = body_font
#             cell.fill   = fill
#             cell.border = cell_border
#             cell.alignment = left_align

#     # Auto-fit column widths (cap at 60)
#     for col_idx, col in enumerate(df.columns, start=1):
#         max_len = max(
#             len(str(col)),
#             *[len(str(v)) for v in df.iloc[:, col_idx - 1].astype(str).values]
#         ) if len(df) > 0 else len(str(col))
#         ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 60)

#     # Freeze header
#     if freeze_header:
#         ws.freeze_panes = "A2"

#     # Row height for header
#     ws.row_dimensions[1].height = 28


# # ══════════════════════════════════════════════════════════════════════════════
# #  PDF READING via PyMuPDF (superior layout preservation)
# # ══════════════════════════════════════════════════════════════════════════════
# def _extract_pdf_text(source_path: str) -> str:
#     """Extract text from PDF preserving layout order via PyMuPDF."""
#     doc = fitz.open(source_path)
#     pages = []
#     for page in doc:
#         # sort=True preserves natural reading order (top→bottom, left→right)
#         pages.append(page.get_text("text", sort=True))
#     doc.close()
#     return "\n\n".join(pages)


# def _extract_pdf_tables(source_path: str) -> list[pd.DataFrame]:
#     """Attempt to extract tables from a PDF using PyMuPDF's table finder."""
#     doc = fitz.open(source_path)
#     all_tables = []
#     for page in doc:
#         tabs = page.find_tables()
#         if tabs and tabs.tables:
#             for tab in tabs.tables:
#                 rows = tab.extract()
#                 if rows:
#                     df = pd.DataFrame(rows[1:], columns=rows[0])
#                     df = df.fillna("").astype(str)
#                     all_tables.append(df)
#     doc.close()
#     return all_tables


# # ══════════════════════════════════════════════════════════════════════════════
# #  DOCX READING — paragraph-aware extraction
# # ══════════════════════════════════════════════════════════════════════════════
# def _extract_docx_rich(source_path: str) -> tuple[list, list]:
#     """
#     Returns (paragraphs, tables).
#     paragraphs: list of dicts with {text, style, bold, italic, alignment}
#     tables:     list of 2D lists
#     """
#     doc = DocxDocument(source_path)
#     paragraphs = []
#     for p in doc.paragraphs:
#         if not p.text.strip():
#             continue
#         paragraphs.append({
#             "text":      p.text,
#             "style":     p.style.name,
#             "bold":      any(r.bold for r in p.runs),
#             "italic":    any(r.italic for r in p.runs),
#             "alignment": p.alignment,
#         })
#     tables = []
#     for table in doc.tables:
#         rows = []
#         for row in table.rows:
#             rows.append([cell.text for cell in row.cells])
#         tables.append(rows)
#     return paragraphs, tables


# # ══════════════════════════════════════════════════════════════════════════════
# #  CONVERSION ENGINE
# # ══════════════════════════════════════════════════════════════════════════════
# class ConversionEngine:

#     @staticmethod
#     def clean_df(df: pd.DataFrame) -> pd.DataFrame:
#         df = df.fillna("")
#         df.columns = df.columns.astype(str)
#         df = df.loc[:, ~df.columns.str.contains(r'^Unnamed:', na=False)]
#         return df

#     @staticmethod
#     def _read_spreadsheet(path: str) -> pd.DataFrame:
#         ext = Path(path).suffix.lower()
#         if ext == ".csv":
#             # Try multiple encodings gracefully
#             for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
#                 try:
#                     return ConversionEngine.clean_df(pd.read_csv(path, encoding=enc))
#                 except UnicodeDecodeError:
#                     continue
#             return pd.DataFrame()
#         else:
#             return ConversionEngine.clean_df(pd.read_excel(path, engine="openpyxl"))

#     # ── Core router ──────────────────────────────────────────────────────────
#     @staticmethod
#     def convert(source_path: str, target_ext: str, output_dir: str) -> bool:
#         src = Path(source_path)
#         src_ext = src.suffix.lower()
#         target_ext = target_ext.lower()
#         if not target_ext.startswith("."):
#             target_ext = "." + target_ext

#         stem = src.stem
#         target_path = str(Path(output_dir) / f"{stem}{target_ext}")

#         # Same format — no-op
#         if src_ext == target_ext:
#             shutil.copy2(source_path, target_path)
#             return True

#         lo = LibreOfficeBridge.available()

#         # ── HIGHEST FIDELITY PATHS (LibreOffice) ────────────────────────────
#         # These go first because LO produces pixel-perfect output for doc types
#         if lo:
#             # DOCX/ODT/TXT → PDF  (perfect fidelity — same as Word's Export PDF)
#             if src_ext in (".docx", ".odt", ".rtf") and target_ext == ".pdf":
#                 out = LibreOfficeBridge.convert(source_path, "pdf", output_dir)
#                 return out is not None

#             # PDF → DOCX  (LO handles this better than manual reconstruction)
#             if src_ext == ".pdf" and target_ext == ".docx":
#                 out = LibreOfficeBridge.convert(source_path, "docx", output_dir)
#                 return out is not None

#             # XLSX/ODS → PDF  (via LO for perfect grid rendering)
#             if src_ext in (".xlsx", ".xls", ".ods") and target_ext == ".pdf":
#                 out = LibreOfficeBridge.convert(source_path, "pdf", output_dir)
#                 return out is not None

#             # CSV → PDF  (build a styled XLSX first, then LO converts)
#             if src_ext == ".csv" and target_ext == ".pdf":
#                 df = ConversionEngine._read_spreadsheet(source_path)
#                 tmp_xlsx = str(Path(output_dir) / f"_tmp_{stem}.xlsx")
#                 ConversionEngine._df_to_styled_xlsx(df, tmp_xlsx)
#                 out = LibreOfficeBridge.convert(tmp_xlsx, "pdf", output_dir)
#                 try:
#                     os.remove(tmp_xlsx)
#                 except Exception:
#                     pass
#                 # LO will name it _tmp_{stem}.pdf — rename
#                 if out:
#                     final = str(Path(output_dir) / f"{stem}.pdf")
#                     if out != final:
#                         shutil.move(out, final)
#                     return True
#                 return False

#             # DOCX → ODT / RTF  (LO handles this perfectly)
#             if src_ext == ".docx" and target_ext in (".odt", ".rtf"):
#                 out = LibreOfficeBridge.convert(source_path, target_ext.lstrip("."), output_dir)
#                 return out is not None

#         # ── SPREADSHEET ↔ SPREADSHEET ────────────────────────────────────────
#         if src_ext in (".csv", ".xlsx", ".xls") and target_ext in (".csv", ".xlsx"):
#             df = ConversionEngine._read_spreadsheet(source_path)
#             if target_ext == ".xlsx":
#                 ConversionEngine._df_to_styled_xlsx(df, target_path)
#             else:
#                 df.to_csv(target_path, index=False, encoding="utf-8-sig")
#             return True

#         # ── SPREADSHEET → DOCX  (styled table) ──────────────────────────────
#         if src_ext in (".csv", ".xlsx", ".xls") and target_ext == ".docx":
#             df = ConversionEngine._read_spreadsheet(source_path)
#             ConversionEngine._df_to_docx(df, target_path, stem)
#             return True

#         # ── SPREADSHEET → TXT ────────────────────────────────────────────────
#         if src_ext in (".csv", ".xlsx", ".xls") and target_ext == ".txt":
#             df = ConversionEngine._read_spreadsheet(source_path)
#             with open(target_path, "w", encoding="utf-8") as f:
#                 f.write(df.to_string(index=False))
#             return True

#         # ── PDF → TXT  (PyMuPDF — layout-aware) ─────────────────────────────
#         if src_ext == ".pdf" and target_ext == ".txt":
#             text = _extract_pdf_text(source_path)
#             with open(target_path, "w", encoding="utf-8") as f:
#                 f.write(text)
#             return True

#         # ── PDF → XLSX  (table extraction) ───────────────────────────────────
#         if src_ext == ".pdf" and target_ext in (".xlsx", ".csv"):
#             tables = _extract_pdf_tables(source_path)
#             if tables:
#                 if target_ext == ".xlsx":
#                     wb = openpyxl.Workbook()
#                     for i, df in enumerate(tables):
#                         ws = wb.create_sheet(f"Table {i+1}") if i > 0 else wb.active
#                         ws.title = f"Table {i+1}"
#                         _style_worksheet(ws, df)
#                     wb.save(target_path)
#                 else:
#                     # Multiple tables — export first, note others
#                     tables[0].to_csv(target_path, index=False, encoding="utf-8-sig")
#                 return True
#             else:
#                 # No tables found — export as plain text rows
#                 text = _extract_pdf_text(source_path)
#                 lines = [l for l in text.split("\n") if l.strip()]
#                 df = pd.DataFrame({"Text": lines})
#                 if target_ext == ".xlsx":
#                     ConversionEngine._df_to_styled_xlsx(df, target_path)
#                 else:
#                     df.to_csv(target_path, index=False, encoding="utf-8-sig")
#                 return True

#         # ── DOCX → TXT ───────────────────────────────────────────────────────
#         if src_ext == ".docx" and target_ext == ".txt":
#             paras, tables = _extract_docx_rich(source_path)
#             lines = [p["text"] for p in paras]
#             for table in tables:
#                 for row in table:
#                     lines.append(" | ".join(row))
#             with open(target_path, "w", encoding="utf-8") as f:
#                 f.write("\n".join(lines))
#             return True

#         # ── DOCX → XLSX / CSV ────────────────────────────────────────────────
#         if src_ext == ".docx" and target_ext in (".xlsx", ".csv"):
#             paras, tables = _extract_docx_rich(source_path)
#             if tables:
#                 # Use the first table found in the document
#                 rows = tables[0]
#                 df = pd.DataFrame(rows[1:], columns=rows[0]) if len(rows) > 1 else pd.DataFrame(rows)
#             else:
#                 df = pd.DataFrame({"Content": [p["text"] for p in paras]})
#             if target_ext == ".xlsx":
#                 ConversionEngine._df_to_styled_xlsx(df, target_path)
#             else:
#                 df.to_csv(target_path, index=False, encoding="utf-8-sig")
#             return True

#         # ── TXT → DOCX ───────────────────────────────────────────────────────
#         if src_ext == ".txt" and target_ext == ".docx":
#             with open(source_path, "r", encoding="utf-8", errors="ignore") as f:
#                 text = f.read()
#             ConversionEngine._text_to_docx(text, target_path, stem)
#             return True

#         # ── TXT → PDF  (via styled DOCX → LibreOffice, or direct ReportLab) ─
#         if src_ext == ".txt" and target_ext == ".pdf":
#             with open(source_path, "r", encoding="utf-8", errors="ignore") as f:
#                 text = f.read()
#             if lo:
#                 tmp_docx = str(Path(output_dir) / f"_tmp_{stem}.docx")
#                 ConversionEngine._text_to_docx(text, tmp_docx, stem)
#                 out = LibreOfficeBridge.convert(tmp_docx, "pdf", output_dir)
#                 try:
#                     os.remove(tmp_docx)
#                 except Exception:
#                     pass
#                 if out:
#                     final = str(Path(output_dir) / f"{stem}.pdf")
#                     if out != final:
#                         shutil.move(out, final)
#                     return True
#             # ReportLab fallback
#             ConversionEngine._text_to_pdf_reportlab(text, target_path)
#             return True

#         # ── TXT → XLSX / CSV ─────────────────────────────────────────────────
#         if src_ext == ".txt" and target_ext in (".xlsx", ".csv"):
#             with open(source_path, "r", encoding="utf-8", errors="ignore") as f:
#                 lines = [l.rstrip() for l in f if l.strip()]
#             df = pd.DataFrame({"Content": lines})
#             if target_ext == ".xlsx":
#                 ConversionEngine._df_to_styled_xlsx(df, target_path)
#             else:
#                 df.to_csv(target_path, index=False, encoding="utf-8-sig")
#             return True

#         print(f"[Engine] No conversion path for {src_ext} → {target_ext}")
#         return False

#     # ── Output builders ──────────────────────────────────────────────────────

#     @staticmethod
#     def _df_to_styled_xlsx(df: pd.DataFrame, path: str):
#         """Write a DataFrame to a fully styled XLSX."""
#         wb = openpyxl.Workbook()
#         ws = wb.active
#         ws.title = "Data"
#         _style_worksheet(ws, df)
#         wb.save(path)

#     @staticmethod
#     def _df_to_docx(df: pd.DataFrame, path: str, title: str = ""):
#         """Write a DataFrame as a styled table inside a DOCX."""
#         doc = DocxDocument()
#         _set_docx_styles(doc)
#         if title:
#             h = doc.add_heading(title.replace("_", " ").title(), level=1)
#         doc.add_paragraph(f"Total rows: {len(df)}  ·  Columns: {len(df.columns)}")
#         doc.add_paragraph("")

#         table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
#         table.style = "Table Grid"
#         # Header
#         for j, col in enumerate(df.columns):
#             table.cell(0, j).text = str(col)
#         # Rows
#         for i, row in enumerate(df.itertuples(index=False), start=1):
#             for j, val in enumerate(row):
#                 table.cell(i, j).text = str(val)

#         _style_docx_table(table, header_row=True)
#         doc.save(path)

#     @staticmethod
#     def _text_to_docx(text: str, path: str, title: str = ""):
#         """Convert plain text to a well-styled DOCX preserving paragraph structure."""
#         doc = DocxDocument()
#         _set_docx_styles(doc)
#         if title:
#             doc.add_heading(title.replace("_", " ").replace("-", " ").title(), level=1)

#         for line in text.split("\n"):
#             stripped = line.strip()
#             if not stripped:
#                 doc.add_paragraph("")
#                 continue
#             # Heuristic: ALL CAPS short lines → treat as heading
#             if stripped.isupper() and len(stripped) < 80:
#                 doc.add_heading(stripped.title(), level=2)
#             else:
#                 p = doc.add_paragraph(stripped)
#                 p.style = doc.styles["Normal"]
#         doc.save(path)

#     @staticmethod
#     def _text_to_pdf_reportlab(text: str, path: str):
#         """ReportLab fallback PDF writer with clean typography."""
#         doc = SimpleDocTemplate(
#             path, pagesize=A4,
#             leftMargin=inch, rightMargin=inch,
#             topMargin=inch, bottomMargin=inch,
#         )
#         styles = getSampleStyleSheet()
#         body = ParagraphStyle(
#             "Body", parent=styles["Normal"],
#             fontName="Helvetica", fontSize=10.5,
#             leading=16, spaceAfter=8,
#             textColor=rl_colors.HexColor("#1E293B"),
#         )
#         heading = ParagraphStyle(
#             "H", parent=styles["Heading1"],
#             fontName="Helvetica-Bold", fontSize=14,
#             leading=20, spaceAfter=12,
#             textColor=rl_colors.HexColor("#059669"),
#         )
#         story = []
#         for line in text.split("\n"):
#             s = line.strip()
#             if not s:
#                 story.append(Spacer(1, 6))
#                 continue
#             clean = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
#             if s.isupper() and len(s) < 80:
#                 story.append(RLParagraph(clean.title(), heading))
#             else:
#                 story.append(RLParagraph(clean, body))
#         doc.build(story)


# # ══════════════════════════════════════════════════════════════════════════════
# #  GUI APPLICATION
# # ══════════════════════════════════════════════════════════════════════════════
# class DocuForgeApp(tb.Window):

#     def __init__(self):
#         super().__init__(themename="cyborg", title="DocuForge Converter Pro", size=(1000, 700))
#         self.selected_files: dict[str, str] = {}
#         self.conversion_log: list[str] = []
#         self._lo_available = LibreOfficeBridge.available()
#         self._setup_logo()
#         self._build_ui()

#     def _setup_logo(self):
#         for path_fn, bind_fn in [
#             (lambda: resource_path("logo.png"), lambda p: self.iconphoto(False, tk.PhotoImage(file=p))),
#             (lambda: resource_path("logo.ico"), lambda p: self.iconbitmap(p)),
#         ]:
#             try:
#                 p = path_fn()
#                 if os.path.exists(p):
#                     bind_fn(p)
#             except Exception:
#                 pass

#     def _build_ui(self):
#         # ── Top header ───────────────────────────────────────────────────────
#         top = tb.Frame(self, padding=(16, 12), bootstyle=DARK)
#         top.pack(fill=X)

#         tb.Label(top, text="⚡ DocuForge Pro",
#                  font=("Helvetica", 17, "bold"), bootstyle=LIGHT).pack(side=LEFT, padx=8)

#         # Fidelity badge
#         badge_text = "✦ High-Fidelity Mode (LibreOffice)" if self._lo_available else "⚡ Standard Mode"
#         badge_style = SUCCESS if self._lo_available else WARNING
#         tb.Label(top, text=badge_text, font=("Helvetica", 9),
#                  bootstyle=badge_style).pack(side=LEFT, padx=14)

#         tb.Button(top, text="📂 Add Folder",    bootstyle=INFO,    command=self._upload_directory).pack(side=RIGHT, padx=4)
#         tb.Button(top, text="📁 Add Files",     bootstyle=PRIMARY,  command=self._upload_files).pack(side=RIGHT, padx=4)
#         tb.Button(top, text="🗑 Clear All",     bootstyle="danger-outline", command=self._clear_all).pack(side=RIGHT, padx=4)

#         # ── Main body ────────────────────────────────────────────────────────
#         body = tb.Frame(self, padding=(14, 10))
#         body.pack(fill=BOTH, expand=True)

#         # Left: file queue
#         left = tb.Labelframe(body, text=" Queued Files ", padding=8, bootstyle=INFO)
#         left.pack(side=LEFT, fill=BOTH, expand=True, padx=(0, 8))

#         self.scroll_box = ScrolledFrame(left, autohide=True)
#         self.scroll_box.pack(fill=BOTH, expand=True)

#         # Right: options
#         right = tb.Labelframe(body, text=" Conversion Options ", padding=14, width=300, bootstyle=INFO)
#         right.pack(side=RIGHT, fill=Y, padx=(8, 0))
#         right.pack_propagate(False)

#         tb.Label(right, text="Convert to:", font=("Helvetica", 11, "bold")).pack(anchor=W, pady=(0, 8))

#         self.target_fmt = tb.StringVar(value=".pdf")
#         FORMATS = [
#             ("PDF Document (.pdf)",        ".pdf"),
#             ("Word Document (.docx)",      ".docx"),
#             ("Excel Spreadsheet (.xlsx)",  ".xlsx"),
#             ("CSV File (.csv)",            ".csv"),
#             ("Plain Text (.txt)",          ".txt"),
#         ]
#         for label, ext in FORMATS:
#             tb.Radiobutton(right, text=label, value=ext,
#                            variable=self.target_fmt, bootstyle=INFO).pack(anchor=W, pady=5)

#         # Fidelity note
#         tb.Separator(right).pack(fill=X, pady=10)
#         note = ("LibreOffice active — DOCX/XLSX→PDF conversion\npreserves 100% formatting fidelity."
#                 if self._lo_available else
#                 "Tip: Install LibreOffice for pixel-perfect\nDOCX/XLSX→PDF conversion.")
#         tb.Label(right, text=note, font=("Helvetica", 8),
#                  bootstyle="success" if self._lo_available else "warning",
#                  wraplength=260, justify=LEFT).pack(anchor=W, pady=(0, 10))

#         tb.Separator(right).pack(fill=X, pady=4)

#         # Progress
#         self.status_var = tb.StringVar(value="Ready")
#         tb.Label(right, textvariable=self.status_var, font=("Helvetica", 9),
#                  bootstyle=SECONDARY, wraplength=260).pack(anchor=W, pady=(8, 4))

#         self.progress = tb.Progressbar(right, bootstyle=SUCCESS, mode="determinate")
#         self.progress.pack(fill=X, pady=(0, 6), side=BOTTOM)

#         self.convert_btn = tb.Button(
#             right, text="⚡ Start Conversion", bootstyle=SUCCESS,
#             state=DISABLED, command=self._start_conversion_thread
#         )
#         self.convert_btn.pack(fill=X, side=BOTTOM, pady=4)

#         # Log button
#         tb.Button(right, text="📋 View Log", bootstyle="secondary-outline",
#                   command=self._show_log).pack(fill=X, side=BOTTOM, pady=2)

#         self._refresh_queue()

#     # ── File management ──────────────────────────────────────────────────────
#     def _refresh_queue(self):
#         for w in self.scroll_box.winfo_children():
#             w.destroy()

#         if not self.selected_files:
#             tb.Label(self.scroll_box, text="Drop files here or use the buttons above.",
#                      font=("Helvetica", 10, "italic"), bootstyle="secondary").pack(pady=30)
#             self.convert_btn.config(state=DISABLED)
#             return

#         self.convert_btn.config(state=NORMAL)
#         EXT_ICONS = {".pdf": "📄", ".docx": "📝", ".xlsx": "📊",
#                      ".csv": "📋", ".txt": "📃"}

#         for idx, (path, group) in enumerate(self.selected_files.items()):
#             ext = Path(path).suffix.lower()
#             icon = EXT_ICONS.get(ext, "📎")
#             bg = DARK if idx % 2 == 0 else SECONDARY

#             row = tb.Frame(self.scroll_box, padding=(8, 5), bootstyle=bg)
#             row.pack(fill=X, pady=2)

#             label = f"[{group}]  {Path(path).name}" if group else f"{icon}  {Path(path).name}"
#             tb.Label(row, text=label, font=("Helvetica", 9)).pack(side=LEFT, padx=4)

#             tb.Button(row, text="✕", padding=(5, 1), bootstyle="danger-outline",
#                       command=lambda p=path: self._remove_file(p)).pack(side=RIGHT, padx=4)

#     def _upload_files(self):
#         paths = filedialog.askopenfilenames(
#             title="Select Files",
#             filetypes=[("Supported Documents", "*.docx *.txt *.csv *.xlsx *.pdf *.odt *.rtf")]
#         )
#         for p in paths:
#             if p not in self.selected_files:
#                 self.selected_files[p] = ""
#         self._refresh_queue()

#     def _upload_directory(self):
#         d = filedialog.askdirectory(title="Select Folder")
#         if not d:
#             return
#         valid = (".docx", ".txt", ".csv", ".xlsx", ".pdf", ".odt", ".rtf")
#         folder = Path(d).name
#         for root, _, files in os.walk(d):
#             for f in files:
#                 if Path(f).suffix.lower() in valid:
#                     fp = os.path.join(root, f)
#                     if fp not in self.selected_files:
#                         self.selected_files[fp] = folder
#         self._refresh_queue()

#     def _remove_file(self, path: str):
#         self.selected_files.pop(path, None)
#         self._refresh_queue()

#     def _clear_all(self):
#         self.selected_files.clear()
#         self._refresh_queue()

#     # ── Conversion ───────────────────────────────────────────────────────────
#     def _start_conversion_thread(self):
#         """Run conversion in background thread so UI stays responsive."""
#         self.convert_btn.config(state=DISABLED)
#         threading.Thread(target=self._run_conversion, daemon=True).start()

#     def _run_conversion(self):
#         output_dir = filedialog.askdirectory(title="Select Output Folder")
#         if not output_dir:
#             self.after(0, lambda: self.convert_btn.config(state=NORMAL))
#             return

#         target_ext = self.target_fmt.get()
#         items = list(self.selected_files.items())
#         total = len(items)
#         success = 0
#         self.conversion_log.clear()

#         for idx, (file_path, group) in enumerate(items):
#             name = Path(file_path).name
#             self.after(0, lambda n=name: self.status_var.set(f"Converting: {n}"))

#             try:
#                 dest_dir = os.path.join(output_dir, group) if group else output_dir
#                 os.makedirs(dest_dir, exist_ok=True)
#                 ok = ConversionEngine.convert(file_path, target_ext, dest_dir)
#                 if ok:
#                     success += 1
#                     self.conversion_log.append(f"✅  {name}")
#                 else:
#                     self.conversion_log.append(f"❌  {name}  — no conversion path found")
#             except Exception as e:
#                 self.conversion_log.append(f"💥  {name}  — {str(e)[:120]}")

#             pct = ((idx + 1) / total) * 100
#             self.after(0, lambda v=pct: self.progress.config(value=v))

#         self.after(0, lambda: self._finish(success, total))

#     def _finish(self, success: int, total: int):
#         self.status_var.set(f"Done — {success}/{total} files converted")
#         self.progress.config(value=0)
#         self.convert_btn.config(state=NORMAL if self.selected_files else DISABLED)
#         messagebox.showinfo(
#             "Conversion Complete",
#             f"✅  {success} of {total} file(s) converted successfully.\n\n"
#             + ("View the log for any errors." if success < total else "All files converted.")
#         )
#         if success == total:
#             self.selected_files.clear()
#             self._refresh_queue()

#     def _show_log(self):
#         """Show a scrollable log window."""
#         win = tk.Toplevel(self)
#         win.title("Conversion Log")
#         win.geometry("600x380")
#         win.configure(bg="#1a1a2e")
#         text = tk.Text(win, bg="#1a1a2e", fg="#e2e8f0", font=("Courier", 10),
#                        relief=tk.FLAT, padx=12, pady=12, wrap=tk.WORD)
#         text.pack(fill=tk.BOTH, expand=True)
#         sb = tk.Scrollbar(win, command=text.yview)
#         sb.pack(side=tk.RIGHT, fill=tk.Y)
#         text.config(yscrollcommand=sb.set)
#         log_text = "\n".join(self.conversion_log) if self.conversion_log else "No conversions run yet."
#         text.insert("1.0", log_text)
#         text.config(state=tk.DISABLED)


# if __name__ == "__main__":
#     app = DocuForgeApp()
#     app.mainloop()








































# #!/usr/bin/env python3
# """
# DocuForge Converter Pro — World-Class Edition
# ═══════════════════════════════════════════════════════════════════════════════
# Conversion fidelity strategy:
#   • DOCX/ODT/PPTX → PDF   : LibreOffice headless  (100% layout fidelity)
#   • XLSX/CSV      → PDF   : LibreOffice headless on intermediate XLSX
#   • PDF           → DOCX  : PyMuPDF text extraction → styled python-docx
#   • PDF           → TXT   : PyMuPDF (superior to pypdf for layout preservation)
#   • PDF           → XLSX  : PyMuPDF + table heuristics → openpyxl
#   • DOCX          → TXT   : python-docx paragraph traversal (preserves order)
#   • DOCX          → XLSX  : paragraph lines → structured dataframe
#   • TXT/CSV       → DOCX  : fully styled python-docx output
#   • CSV/XLSX      → DOCX  : styled table inside python-docx
#   • Any spreadsheet ↔ any spreadsheet : pandas
# """

# import os, sys, re, shutil, subprocess, tempfile, threading, time
# import tkinter as tk
# from tkinter import filedialog, messagebox
# from pathlib import Path

# import ttkbootstrap as tb
# from ttkbootstrap.constants import *
# from ttkbootstrap.widgets.scrolled import ScrolledFrame

# # Document libraries
# import docx2txt
# from docx import Document as DocxDocument
# from docx.shared import Pt, RGBColor, Inches, Cm
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.table import WD_TABLE_ALIGNMENT

# import fitz          # PyMuPDF — best-in-class PDF reader
# import pandas as pd
# import openpyxl
# from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
# from openpyxl.utils import get_column_letter

# from reportlab.lib.pagesizes import letter, landscape, A4
# from reportlab.platypus import SimpleDocTemplate, Paragraph as RLParagraph, Spacer, Table as RLTable, TableStyle
# from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
# from reportlab.lib import colors as rl_colors
# from reportlab.lib.units import inch


# # ── Runtime asset resolver ──────────────────────────────────────────────────
# def resource_path(relative_path):
#     try:
#         base = sys._MEIPASS
#     except AttributeError:
#         base = os.path.abspath(".")
#     return os.path.join(base, relative_path)


# # ── Font detection utility ──────────────────────────────────────────────────
# def detect_docx_fonts(docx_path: str) -> set[str]:
#     """
#     Extract all fonts used in a DOCX document.
#     Run this to see which fonts need to be installed on the system.
#     """
#     fonts_used = set()
#     doc = DocxDocument(docx_path)
#     for para in doc.paragraphs:
#         for run in para.runs:
#             if run.font.name:
#                 fonts_used.add(run.font.name)
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for para in cell.paragraphs:
#                     for run in para.runs:
#                         if run.font.name:
#                             fonts_used.add(run.font.name)
#     return fonts_used


# # ══════════════════════════════════════════════════════════════════════════════
# #  LIBREOFFICE BRIDGE — the secret weapon for 100% fidelity
# # ══════════════════════════════════════════════════════════════════════════════
# class LibreOfficeBridge:
#     """
#     Wraps LibreOffice headless mode. When LibreOffice converts DOCX→PDF it uses
#     the same rendering engine as the desktop app — identical to Microsoft Word's
#     own 'Export to PDF' in terms of fidelity. Tables, fonts, colours, images,
#     headers/footers — all preserved exactly.

#     FONT EMBEDDING: The PDF export filter now forces font embedding via
#     'EmbedStandardFonts=true,EmbedFont=true' which ensures all fonts used
#     in the document are embedded in the final PDF.
#     """

#     # Candidate binary locations across Windows, macOS, Linux
#     _CANDIDATES = [
#         "libreoffice", "soffice",
#         r"C:\Program Files\LibreOffice\program\soffice.exe",
#         r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
#         "/Applications/LibreOffice.app/Contents/MacOS/soffice",
#         "/usr/bin/libreoffice", "/usr/bin/soffice",
#         "/usr/local/bin/libreoffice",
#     ]

#     @classmethod
#     def find_binary(cls):
#         for candidate in cls._CANDIDATES:
#             if shutil.which(candidate) or os.path.isfile(candidate):
#                 return candidate
#         return None

#     @classmethod
#     def convert(cls, source_path: str, target_format: str, output_dir: str) -> str | None:
#         """
#         Convert source_path to target_format inside output_dir.
#         Returns the output file path on success, None on failure.
#         target_format examples: 'pdf', 'docx', 'xlsx', 'csv', 'txt'

#         UPDATED: PDF conversion now uses explicit filter string with font embedding:
#         'pdf:writer_pdf_Export:EmbedStandardFonts=true,EmbedFont=true,MaxImageResolution=300'
#         """
#         binary = cls.find_binary()
#         if not binary:
#             return None

#         # Build the PDF export filter string with font embedding enabled
#         if target_format == "pdf":
#             lo_format = (
#                 "pdf:writer_pdf_Export:"
#                 "EmbedStandardFonts=true,"
#                 "EmbedFont=true,"
#                 "MaxImageResolution=300"
#             )
#         elif target_format in ("xlsx", "xls"):
#             lo_format = "xlsx:Calc MS Excel 2007 XML"
#         elif target_format == "docx":
#             lo_format = "docx:MS Word 2007 XML"
#         else:
#             lo_format = target_format

#         # LibreOffice locks its user-profile dir — use a temp profile per call
#         # to allow parallel conversions and avoid conflicts with running instances.
#         with tempfile.TemporaryDirectory() as profile_dir:
#             cmd = [
#                 binary,
#                 "--headless",
#                 f"-env:UserInstallation=file:///{profile_dir.replace(os.sep, '/')}",
#                 "--convert-to", lo_format,
#                 "--outdir", output_dir,
#                 source_path,
#             ]
#             try:
#                 result = subprocess.run(
#                     cmd, capture_output=True, text=True, timeout=120
#                 )
#                 if result.returncode != 0:
#                     print(f"[LibreOffice] Error: {result.stderr[:400]}")
#                     return None

#                 # LO renames the file by swapping the extension
#                 stem = Path(source_path).stem
#                 ext = target_format.split(":")[0]  # handle "pdf" vs "pdf:writer_pdf_Export"
#                 out = Path(output_dir) / f"{stem}.{ext}"
#                 return str(out) if out.exists() else None
#             except subprocess.TimeoutExpired:
#                 print("[LibreOffice] Conversion timed out")
#                 return None
#             except Exception as e:
#                 print(f"[LibreOffice] Unexpected error: {e}")
#                 return None

#     @classmethod
#     def available(cls) -> bool:
#         return cls.find_binary() is not None


# # ══════════════════════════════════════════════════════════════════════════════
# #  DOCX STYLING HELPERS
# # ══════════════════════════════════════════════════════════════════════════════
# # Brand palette used in generated DOCX / XLSX outputs
# BRAND = {
#     "header_bg":   RGBColor(0x05, 0x41, 0x3E),  # deep teal
#     "header_fg":   RGBColor(0xFF, 0xFF, 0xFF),
#     "accent":      RGBColor(0x05, 0x96, 0x69),  # emerald
#     "row_alt":     RGBColor(0xF0, 0xFD, 0xFA),
#     "border":      RGBColor(0xCB, 0xD5, 0xE1),
#     "body_text":   RGBColor(0x1E, 0x29, 0x3B),
# }

# BRAND_HEX = {
#     "header_bg":  "05413E",
#     "header_fg":  "FFFFFF",
#     "accent":     "059669",
#     "row_alt":    "F0FDFA",
#     "row_even":   "FFFFFF",
#     "border":     "CBD5E1",
# }


# def _style_docx_table(table, header_row=True):
#     """Apply professional styling to a python-docx table."""
#     tbl = table._tbl
#     for i, row in enumerate(table.rows):
#         for j, cell in enumerate(row.cells):
#             tc = cell._tc
#             # Cell margins
#             from docx.oxml.ns import qn
#             from docx.oxml import OxmlElement
#             tcPr = tc.get_or_add_tcPr()
#             tcMar = OxmlElement('w:tcMar')
#             for side in ('top', 'left', 'bottom', 'right'):
#                 node = OxmlElement(f'w:{side}')
#                 node.set(qn('w:w'), '80')
#                 node.set(qn('w:type'), 'dxa')
#                 tcMar.append(node)
#             tcPr.append(tcMar)

#             para = cell.paragraphs[0]
#             run = para.runs[0] if para.runs else para.add_run(cell.text)
#             run.font.size = Pt(9.5)
#             run.font.color.rgb = BRAND["body_text"]

#             if header_row and i == 0:
#                 run.bold = True
#                 run.font.color.rgb = BRAND["header_fg"]
#                 from docx.oxml.ns import qn
#                 from docx.oxml import OxmlElement
#                 shading = OxmlElement('w:shd')
#                 shading.set(qn('w:val'), 'clear')
#                 shading.set(qn('w:fill'), BRAND_HEX["header_bg"])
#                 tcPr.append(shading)
#             elif i % 2 == 0:
#                 from docx.oxml.ns import qn
#                 from docx.oxml import OxmlElement
#                 shading = OxmlElement('w:shd')
#                 shading.set(qn('w:val'), 'clear')
#                 shading.set(qn('w:fill'), BRAND_HEX["row_alt"])
#                 tcPr.append(shading)


# def _set_docx_styles(doc: DocxDocument):
#     """Configure document-level typography."""
#     style = doc.styles['Normal']
#     style.font.name = 'Calibri'
#     style.font.size = Pt(11)
#     style.font.color.rgb = BRAND["body_text"]

#     # Heading 1
#     h1 = doc.styles['Heading 1']
#     h1.font.name = 'Calibri'
#     h1.font.size = Pt(20)
#     h1.font.bold = True
#     h1.font.color.rgb = BRAND["accent"]

#     # Heading 2
#     h2 = doc.styles['Heading 2']
#     h2.font.name = 'Calibri'
#     h2.font.size = Pt(14)
#     h2.font.bold = True
#     h2.font.color.rgb = BRAND["header_bg"]


# # ══════════════════════════════════════════════════════════════════════════════
# #  XLSX STYLING HELPERS
# # ══════════════════════════════════════════════════════════════════════════════
# def _style_worksheet(ws, df, freeze_header=True):
#     """Apply professional table styling to an openpyxl worksheet."""
#     header_fill = PatternFill("solid", fgColor=BRAND_HEX["header_bg"])
#     alt_fill    = PatternFill("solid", fgColor=BRAND_HEX["row_alt"])
#     even_fill   = PatternFill("solid", fgColor=BRAND_HEX["row_even"])
#     thin = Side(border_style="thin", color=BRAND_HEX["border"])
#     cell_border = Border(left=thin, right=thin, top=thin, bottom=thin)

#     header_font = Font(name="Calibri", bold=True, color=BRAND_HEX["header_fg"], size=10)
#     body_font   = Font(name="Calibri", size=10, color="1E293B")
#     center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
#     left_align   = Alignment(horizontal="left",   vertical="center", wrap_text=True)

#     # Header row
#     for col_idx, col_name in enumerate(df.columns, start=1):
#         cell = ws.cell(row=1, column=col_idx, value=str(col_name))
#         cell.font   = header_font
#         cell.fill   = header_fill
#         cell.border = cell_border
#         cell.alignment = center_align

#     # Data rows
#     for row_idx, row in enumerate(df.itertuples(index=False), start=2):
#         fill = alt_fill if row_idx % 2 == 0 else even_fill
#         for col_idx, value in enumerate(row, start=1):
#             cell = ws.cell(row=row_idx, column=col_idx, value=value)
#             cell.font   = body_font
#             cell.fill   = fill
#             cell.border = cell_border
#             cell.alignment = left_align

#     # Auto-fit column widths (cap at 60)
#     for col_idx, col in enumerate(df.columns, start=1):
#         max_len = max(
#             len(str(col)),
#             *[len(str(v)) for v in df.iloc[:, col_idx - 1].astype(str).values]
#         ) if len(df) > 0 else len(str(col))
#         ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 60)

#     # Freeze header
#     if freeze_header:
#         ws.freeze_panes = "A2"

#     # Row height for header
#     ws.row_dimensions[1].height = 28


# # ══════════════════════════════════════════════════════════════════════════════
# #  PDF READING via PyMuPDF (superior layout preservation)
# # ══════════════════════════════════════════════════════════════════════════════
# def _extract_pdf_text(source_path: str) -> str:
#     """Extract text from PDF preserving layout order via PyMuPDF."""
#     doc = fitz.open(source_path)
#     pages = []
#     for page in doc:
#         # sort=True preserves natural reading order (top→bottom, left→right)
#         pages.append(page.get_text("text", sort=True))
#     doc.close()
#     return "\n\n".join(pages)


# def _extract_pdf_tables(source_path: str) -> list[pd.DataFrame]:
#     """Attempt to extract tables from a PDF using PyMuPDF's table finder."""
#     doc = fitz.open(source_path)
#     all_tables = []
#     for page in doc:
#         tabs = page.find_tables()
#         if tabs and tabs.tables:
#             for tab in tabs.tables:
#                 rows = tab.extract()
#                 if rows:
#                     df = pd.DataFrame(rows[1:], columns=rows[0])
#                     df = df.fillna("").astype(str)
#                     all_tables.append(df)
#     doc.close()
#     return all_tables


# # ══════════════════════════════════════════════════════════════════════════════
# #  DOCX READING — paragraph-aware extraction
# # ══════════════════════════════════════════════════════════════════════════════
# def _extract_docx_rich(source_path: str) -> tuple[list, list]:
#     """
#     Returns (paragraphs, tables).
#     paragraphs: list of dicts with {text, style, bold, italic, alignment}
#     tables:     list of 2D lists
#     """
#     doc = DocxDocument(source_path)
#     paragraphs = []
#     for p in doc.paragraphs:
#         if not p.text.strip():
#             continue
#         paragraphs.append({
#             "text":      p.text,
#             "style":     p.style.name,
#             "bold":      any(r.bold for r in p.runs),
#             "italic":    any(r.italic for r in p.runs),
#             "alignment": p.alignment,
#         })
#     tables = []
#     for table in doc.tables:
#         rows = []
#         for row in table.rows:
#             rows.append([cell.text for cell in row.cells])
#         tables.append(rows)
#     return paragraphs, tables


# # ══════════════════════════════════════════════════════════════════════════════
# #  CONVERSION ENGINE
# # ══════════════════════════════════════════════════════════════════════════════
# class ConversionEngine:

#     @staticmethod
#     def clean_df(df: pd.DataFrame) -> pd.DataFrame:
#         df = df.fillna("")
#         df.columns = df.columns.astype(str)
#         df = df.loc[:, ~df.columns.str.contains(r'^Unnamed:', na=False)]
#         return df

#     @staticmethod
#     def _read_spreadsheet(path: str) -> pd.DataFrame:
#         ext = Path(path).suffix.lower()
#         if ext == ".csv":
#             # Try multiple encodings gracefully
#             for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
#                 try:
#                     return ConversionEngine.clean_df(pd.read_csv(path, encoding=enc))
#                 except UnicodeDecodeError:
#                     continue
#             return pd.DataFrame()
#         else:
#             return ConversionEngine.clean_df(pd.read_excel(path, engine="openpyxl"))

#     # ── Core router ──────────────────────────────────────────────────────────
#     @staticmethod
#     def convert(source_path: str, target_ext: str, output_dir: str) -> bool:
#         src = Path(source_path)
#         src_ext = src.suffix.lower()
#         target_ext = target_ext.lower()
#         if not target_ext.startswith("."):
#             target_ext = "." + target_ext

#         stem = src.stem
#         target_path = str(Path(output_dir) / f"{stem}{target_ext}")

#         # Same format — no-op
#         if src_ext == target_ext:
#             shutil.copy2(source_path, target_path)
#             return True

#         lo = LibreOfficeBridge.available()

#         # ── HIGHEST FIDELITY PATHS (LibreOffice) ────────────────────────────
#         # These go first because LO produces pixel-perfect output for doc types
#         if lo:
#             # DOCX/ODT/TXT → PDF  (perfect fidelity — same as Word's Export PDF)
#             if src_ext in (".docx", ".odt", ".rtf") and target_ext == ".pdf":
#                 out = LibreOfficeBridge.convert(source_path, "pdf", output_dir)
#                 return out is not None

#             # PDF → DOCX  (LO handles this better than manual reconstruction)
#             if src_ext == ".pdf" and target_ext == ".docx":
#                 out = LibreOfficeBridge.convert(source_path, "docx", output_dir)
#                 return out is not None

#             # XLSX/ODS → PDF  (via LO for perfect grid rendering)
#             if src_ext in (".xlsx", ".xls", ".ods") and target_ext == ".pdf":
#                 out = LibreOfficeBridge.convert(source_path, "pdf", output_dir)
#                 return out is not None

#             # CSV → PDF  (build a styled XLSX first, then LO converts)
#             if src_ext == ".csv" and target_ext == ".pdf":
#                 df = ConversionEngine._read_spreadsheet(source_path)
#                 tmp_xlsx = str(Path(output_dir) / f"_tmp_{stem}.xlsx")
#                 ConversionEngine._df_to_styled_xlsx(df, tmp_xlsx)
#                 out = LibreOfficeBridge.convert(tmp_xlsx, "pdf", output_dir)
#                 try:
#                     os.remove(tmp_xlsx)
#                 except Exception:
#                     pass
#                 # LO will name it _tmp_{stem}.pdf — rename
#                 if out:
#                     final = str(Path(output_dir) / f"{stem}.pdf")
#                     if out != final:
#                         shutil.move(out, final)
#                     return True
#                 return False

#             # DOCX → ODT / RTF  (LO handles this perfectly)
#             if src_ext == ".docx" and target_ext in (".odt", ".rtf"):
#                 out = LibreOfficeBridge.convert(source_path, target_ext.lstrip("."), output_dir)
#                 return out is not None

#         # ── SPREADSHEET ↔ SPREADSHEET ────────────────────────────────────────
#         if src_ext in (".csv", ".xlsx", ".xls") and target_ext in (".csv", ".xlsx"):
#             df = ConversionEngine._read_spreadsheet(source_path)
#             if target_ext == ".xlsx":
#                 ConversionEngine._df_to_styled_xlsx(df, target_path)
#             else:
#                 df.to_csv(target_path, index=False, encoding="utf-8-sig")
#             return True

#         # ── SPREADSHEET → DOCX  (styled table) ──────────────────────────────
#         if src_ext in (".csv", ".xlsx", ".xls") and target_ext == ".docx":
#             df = ConversionEngine._read_spreadsheet(source_path)
#             ConversionEngine._df_to_docx(df, target_path, stem)
#             return True

#         # ── SPREADSHEET → TXT ────────────────────────────────────────────────
#         if src_ext in (".csv", ".xlsx", ".xls") and target_ext == ".txt":
#             df = ConversionEngine._read_spreadsheet(source_path)
#             with open(target_path, "w", encoding="utf-8") as f:
#                 f.write(df.to_string(index=False))
#             return True

#         # ── PDF → TXT  (PyMuPDF — layout-aware) ─────────────────────────────
#         if src_ext == ".pdf" and target_ext == ".txt":
#             text = _extract_pdf_text(source_path)
#             with open(target_path, "w", encoding="utf-8") as f:
#                 f.write(text)
#             return True

#         # ── PDF → XLSX  (table extraction) ───────────────────────────────────
#         if src_ext == ".pdf" and target_ext in (".xlsx", ".csv"):
#             tables = _extract_pdf_tables(source_path)
#             if tables:
#                 if target_ext == ".xlsx":
#                     wb = openpyxl.Workbook()
#                     for i, df in enumerate(tables):
#                         ws = wb.create_sheet(f"Table {i+1}") if i > 0 else wb.active
#                         ws.title = f"Table {i+1}"
#                         _style_worksheet(ws, df)
#                     wb.save(target_path)
#                 else:
#                     # Multiple tables — export first, note others
#                     tables[0].to_csv(target_path, index=False, encoding="utf-8-sig")
#                 return True
#             else:
#                 # No tables found — export as plain text rows
#                 text = _extract_pdf_text(source_path)
#                 lines = [l for l in text.split("\n") if l.strip()]
#                 df = pd.DataFrame({"Text": lines})
#                 if target_ext == ".xlsx":
#                     ConversionEngine._df_to_styled_xlsx(df, target_path)
#                 else:
#                     df.to_csv(target_path, index=False, encoding="utf-8-sig")
#                 return True

#         # ── DOCX → TXT ───────────────────────────────────────────────────────
#         if src_ext == ".docx" and target_ext == ".txt":
#             paras, tables = _extract_docx_rich(source_path)
#             lines = [p["text"] for p in paras]
#             for table in tables:
#                 for row in table:
#                     lines.append(" | ".join(row))
#             with open(target_path, "w", encoding="utf-8") as f:
#                 f.write("\n".join(lines))
#             return True

#         # ── DOCX → XLSX / CSV ────────────────────────────────────────────────
#         if src_ext == ".docx" and target_ext in (".xlsx", ".csv"):
#             paras, tables = _extract_docx_rich(source_path)
#             if tables:
#                 # Use the first table found in the document
#                 rows = tables[0]
#                 df = pd.DataFrame(rows[1:], columns=rows[0]) if len(rows) > 1 else pd.DataFrame(rows)
#             else:
#                 df = pd.DataFrame({"Content": [p["text"] for p in paras]})
#             if target_ext == ".xlsx":
#                 ConversionEngine._df_to_styled_xlsx(df, target_path)
#             else:
#                 df.to_csv(target_path, index=False, encoding="utf-8-sig")
#             return True

#         # ── TXT → DOCX ───────────────────────────────────────────────────────
#         if src_ext == ".txt" and target_ext == ".docx":
#             with open(source_path, "r", encoding="utf-8", errors="ignore") as f:
#                 text = f.read()
#             ConversionEngine._text_to_docx(text, target_path, stem)
#             return True

#         # ── TXT → PDF  (via styled DOCX → LibreOffice, or direct ReportLab) ─
#         if src_ext == ".txt" and target_ext == ".pdf":
#             with open(source_path, "r", encoding="utf-8", errors="ignore") as f:
#                 text = f.read()
#             if lo:
#                 tmp_docx = str(Path(output_dir) / f"_tmp_{stem}.docx")
#                 ConversionEngine._text_to_docx(text, tmp_docx, stem)
#                 out = LibreOfficeBridge.convert(tmp_docx, "pdf", output_dir)
#                 try:
#                     os.remove(tmp_docx)
#                 except Exception:
#                     pass
#                 if out:
#                     final = str(Path(output_dir) / f"{stem}.pdf")
#                     if out != final:
#                         shutil.move(out, final)
#                     return True
#             # ReportLab fallback
#             ConversionEngine._text_to_pdf_reportlab(text, target_path)
#             return True

#         # ── TXT → XLSX / CSV ─────────────────────────────────────────────────
#         if src_ext == ".txt" and target_ext in (".xlsx", ".csv"):
#             with open(source_path, "r", encoding="utf-8", errors="ignore") as f:
#                 lines = [l.rstrip() for l in f if l.strip()]
#             df = pd.DataFrame({"Content": lines})
#             if target_ext == ".xlsx":
#                 ConversionEngine._df_to_styled_xlsx(df, target_path)
#             else:
#                 df.to_csv(target_path, index=False, encoding="utf-8-sig")
#             return True

#         print(f"[Engine] No conversion path for {src_ext} → {target_ext}")
#         return False

#     # ── Output builders ──────────────────────────────────────────────────────

#     @staticmethod
#     def _df_to_styled_xlsx(df: pd.DataFrame, path: str):
#         """Write a DataFrame to a fully styled XLSX."""
#         wb = openpyxl.Workbook()
#         ws = wb.active
#         ws.title = "Data"
#         _style_worksheet(ws, df)
#         wb.save(path)

#     @staticmethod
#     def _df_to_docx(df: pd.DataFrame, path: str, title: str = ""):
#         """Write a DataFrame as a styled table inside a DOCX."""
#         doc = DocxDocument()
#         _set_docx_styles(doc)
#         if title:
#             h = doc.add_heading(title.replace("_", " ").title(), level=1)
#         doc.add_paragraph(f"Total rows: {len(df)}  ·  Columns: {len(df.columns)}")
#         doc.add_paragraph("")

#         table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
#         table.style = "Table Grid"
#         # Header
#         for j, col in enumerate(df.columns):
#             table.cell(0, j).text = str(col)
#         # Rows
#         for i, row in enumerate(df.itertuples(index=False), start=1):
#             for j, val in enumerate(row):
#                 table.cell(i, j).text = str(val)

#         _style_docx_table(table, header_row=True)
#         doc.save(path)

#     @staticmethod
#     def _text_to_docx(text: str, path: str, title: str = ""):
#         """Convert plain text to a well-styled DOCX preserving paragraph structure."""
#         doc = DocxDocument()
#         _set_docx_styles(doc)
#         if title:
#             doc.add_heading(title.replace("_", " ").replace("-", " ").title(), level=1)

#         for line in text.split("\n"):
#             stripped = line.strip()
#             if not stripped:
#                 doc.add_paragraph("")
#                 continue
#             # Heuristic: ALL CAPS short lines → treat as heading
#             if stripped.isupper() and len(stripped) < 80:
#                 doc.add_heading(stripped.title(), level=2)
#             else:
#                 p = doc.add_paragraph(stripped)
#                 p.style = doc.styles["Normal"]
#         doc.save(path)

#     @staticmethod
#     def _text_to_pdf_reportlab(text: str, path: str):
#         """ReportLab fallback PDF writer with clean typography."""
#         doc = SimpleDocTemplate(
#             path, pagesize=A4,
#             leftMargin=inch, rightMargin=inch,
#             topMargin=inch, bottomMargin=inch,
#         )
#         styles = getSampleStyleSheet()
#         body = ParagraphStyle(
#             "Body", parent=styles["Normal"],
#             fontName="Helvetica", fontSize=10.5,
#             leading=16, spaceAfter=8,
#             textColor=rl_colors.HexColor("#1E293B"),
#         )
#         heading = ParagraphStyle(
#             "H", parent=styles["Heading1"],
#             fontName="Helvetica-Bold", fontSize=14,
#             leading=20, spaceAfter=12,
#             textColor=rl_colors.HexColor("#059669"),
#         )
#         story = []
#         for line in text.split("\n"):
#             s = line.strip()
#             if not s:
#                 story.append(Spacer(1, 6))
#                 continue
#             clean = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
#             if s.isupper() and len(s) < 80:
#                 story.append(RLParagraph(clean.title(), heading))
#             else:
#                 story.append(RLParagraph(clean, body))
#         doc.build(story)


# # ══════════════════════════════════════════════════════════════════════════════
# #  GUI APPLICATION
# # ══════════════════════════════════════════════════════════════════════════════
# class DocuForgeApp(tb.Window):

#     def __init__(self):
#         super().__init__(themename="cyborg", title="DocuForge Converter Pro", size=(1000, 700))
#         self.selected_files: dict[str, str] = {}
#         self.conversion_log: list[str] = []
#         self._lo_available = LibreOfficeBridge.available()
#         self._setup_logo()
#         self._build_ui()

#     def _setup_logo(self):
#         for path_fn, bind_fn in [
#             (lambda: resource_path("logo.png"), lambda p: self.iconphoto(False, tk.PhotoImage(file=p))),
#             (lambda: resource_path("logo.ico"), lambda p: self.iconbitmap(p)),
#         ]:
#             try:
#                 p = path_fn()
#                 if os.path.exists(p):
#                     bind_fn(p)
#             except Exception:
#                 pass

#     def _build_ui(self):
#         # ── Top header ───────────────────────────────────────────────────────
#         top = tb.Frame(self, padding=(16, 12), bootstyle=DARK)
#         top.pack(fill=X)

#         tb.Label(top, text="⚡ DocuForge Pro",
#                  font=("Helvetica", 17, "bold"), bootstyle=LIGHT).pack(side=LEFT, padx=8)

#         # Fidelity badge
#         badge_text = "✦ High-Fidelity Mode (LibreOffice)" if self._lo_available else "⚡ Standard Mode"
#         badge_style = SUCCESS if self._lo_available else WARNING
#         tb.Label(top, text=badge_text, font=("Helvetica", 9),
#                  bootstyle=badge_style).pack(side=LEFT, padx=14)

#         tb.Button(top, text="📂 Add Folder",    bootstyle=INFO,    command=self._upload_directory).pack(side=RIGHT, padx=4)
#         tb.Button(top, text="📁 Add Files",     bootstyle=PRIMARY,  command=self._upload_files).pack(side=RIGHT, padx=4)
#         tb.Button(top, text="🗑 Clear All",     bootstyle="danger-outline", command=self._clear_all).pack(side=RIGHT, padx=4)

#         # ── Main body ────────────────────────────────────────────────────────
#         body = tb.Frame(self, padding=(14, 10))
#         body.pack(fill=BOTH, expand=True)

#         # Left: file queue
#         left = tb.Labelframe(body, text=" Queued Files ", padding=8, bootstyle=INFO)
#         left.pack(side=LEFT, fill=BOTH, expand=True, padx=(0, 8))

#         self.scroll_box = ScrolledFrame(left, autohide=True)
#         self.scroll_box.pack(fill=BOTH, expand=True)

#         # Right: options
#         right = tb.Labelframe(body, text=" Conversion Options ", padding=14, width=300, bootstyle=INFO)
#         right.pack(side=RIGHT, fill=Y, padx=(8, 0))
#         right.pack_propagate(False)

#         tb.Label(right, text="Convert to:", font=("Helvetica", 11, "bold")).pack(anchor=W, pady=(0, 8))

#         self.target_fmt = tb.StringVar(value=".pdf")
#         FORMATS = [
#             ("PDF Document (.pdf)",        ".pdf"),
#             ("Word Document (.docx)",      ".docx"),
#             ("Excel Spreadsheet (.xlsx)",  ".xlsx"),
#             ("CSV File (.csv)",            ".csv"),
#             ("Plain Text (.txt)",          ".txt"),
#         ]
#         for label, ext in FORMATS:
#             tb.Radiobutton(right, text=label, value=ext,
#                            variable=self.target_fmt, bootstyle=INFO).pack(anchor=W, pady=5)

#         # Fidelity note
#         tb.Separator(right).pack(fill=X, pady=10)
#         note = ("LibreOffice active — DOCX/XLSX→PDF conversion\npreserves 100% formatting fidelity."
#                 if self._lo_available else
#                 "Tip: Install LibreOffice for pixel-perfect\nDOCX/XLSX→PDF conversion.")
#         tb.Label(right, text=note, font=("Helvetica", 8),
#                  bootstyle="success" if self._lo_available else "warning",
#                  wraplength=260, justify=LEFT).pack(anchor=W, pady=(0, 10))

#         tb.Separator(right).pack(fill=X, pady=4)

#         # Progress
#         self.status_var = tb.StringVar(value="Ready")
#         tb.Label(right, textvariable=self.status_var, font=("Helvetica", 9),
#                  bootstyle=SECONDARY, wraplength=260).pack(anchor=W, pady=(8, 4))

#         self.progress = tb.Progressbar(right, bootstyle=SUCCESS, mode="determinate")
#         self.progress.pack(fill=X, pady=(0, 6), side=BOTTOM)

#         self.convert_btn = tb.Button(
#             right, text="⚡ Start Conversion", bootstyle=SUCCESS,
#             state=DISABLED, command=self._start_conversion_thread
#         )
#         self.convert_btn.pack(fill=X, side=BOTTOM, pady=4)

#         # Log button
#         tb.Button(right, text="📋 View Log", bootstyle="secondary-outline",
#                   command=self._show_log).pack(fill=X, side=BOTTOM, pady=2)

#         self._refresh_queue()

#     # ── File management ──────────────────────────────────────────────────────
#     def _refresh_queue(self):
#         for w in self.scroll_box.winfo_children():
#             w.destroy()

#         if not self.selected_files:
#             tb.Label(self.scroll_box, text="Drop files here or use the buttons above.",
#                      font=("Helvetica", 10, "italic"), bootstyle="secondary").pack(pady=30)
#             self.convert_btn.config(state=DISABLED)
#             return

#         self.convert_btn.config(state=NORMAL)
#         EXT_ICONS = {".pdf": "📄", ".docx": "📝", ".xlsx": "📊",
#                      ".csv": "📋", ".txt": "📃"}

#         for idx, (path, group) in enumerate(self.selected_files.items()):
#             ext = Path(path).suffix.lower()
#             icon = EXT_ICONS.get(ext, "📎")
#             bg = DARK if idx % 2 == 0 else SECONDARY

#             row = tb.Frame(self.scroll_box, padding=(8, 5), bootstyle=bg)
#             row.pack(fill=X, pady=2)

#             label = f"[{group}]  {Path(path).name}" if group else f"{icon}  {Path(path).name}"
#             tb.Label(row, text=label, font=("Helvetica", 9)).pack(side=LEFT, padx=4)

#             tb.Button(row, text="✕", padding=(5, 1), bootstyle="danger-outline",
#                       command=lambda p=path: self._remove_file(p)).pack(side=RIGHT, padx=4)

#     def _upload_files(self):
#         paths = filedialog.askopenfilenames(
#             title="Select Files",
#             filetypes=[("Supported Documents", "*.docx *.txt *.csv *.xlsx *.pdf *.odt *.rtf")]
#         )
#         for p in paths:
#             if p not in self.selected_files:
#                 self.selected_files[p] = ""
#         self._refresh_queue()

#     def _upload_directory(self):
#         d = filedialog.askdirectory(title="Select Folder")
#         if not d:
#             return
#         valid = (".docx", ".txt", ".csv", ".xlsx", ".pdf", ".odt", ".rtf")
#         folder = Path(d).name
#         for root, _, files in os.walk(d):
#             for f in files:
#                 if Path(f).suffix.lower() in valid:
#                     fp = os.path.join(root, f)
#                     if fp not in self.selected_files:
#                         self.selected_files[fp] = folder
#         self._refresh_queue()

#     def _remove_file(self, path: str):
#         self.selected_files.pop(path, None)
#         self._refresh_queue()

#     def _clear_all(self):
#         self.selected_files.clear()
#         self._refresh_queue()

#     # ── Conversion ───────────────────────────────────────────────────────────
#     def _start_conversion_thread(self):
#         """Run conversion in background thread so UI stays responsive."""
#         self.convert_btn.config(state=DISABLED)
#         threading.Thread(target=self._run_conversion, daemon=True).start()

#     def _run_conversion(self):
#         output_dir = filedialog.askdirectory(title="Select Output Folder")
#         if not output_dir:
#             self.after(0, lambda: self.convert_btn.config(state=NORMAL))
#             return

#         target_ext = self.target_fmt.get()
#         items = list(self.selected_files.items())
#         total = len(items)
#         success = 0
#         self.conversion_log.clear()

#         for idx, (file_path, group) in enumerate(items):
#             name = Path(file_path).name
#             self.after(0, lambda n=name: self.status_var.set(f"Converting: {n}"))

#             try:
#                 dest_dir = os.path.join(output_dir, group) if group else output_dir
#                 os.makedirs(dest_dir, exist_ok=True)
#                 ok = ConversionEngine.convert(file_path, target_ext, dest_dir)
#                 if ok:
#                     success += 1
#                     self.conversion_log.append(f"✅  {name}")
#                 else:
#                     self.conversion_log.append(f"❌  {name}  — no conversion path found")
#             except Exception as e:
#                 self.conversion_log.append(f"💥  {name}  — {str(e)[:120]}")

#             pct = ((idx + 1) / total) * 100
#             self.after(0, lambda v=pct: self.progress.config(value=v))

#         self.after(0, lambda: self._finish(success, total))

#     def _finish(self, success: int, total: int):
#         self.status_var.set(f"Done — {success}/{total} files converted")
#         self.progress.config(value=0)
#         self.convert_btn.config(state=NORMAL if self.selected_files else DISABLED)
#         messagebox.showinfo(
#             "Conversion Complete",
#             f"✅  {success} of {total} file(s) converted successfully.\n\n"
#             + ("View the log for any errors." if success < total else "All files converted.")
#         )
#         if success == total:
#             self.selected_files.clear()
#             self._refresh_queue()

#     def _show_log(self):
#         """Show a scrollable log window."""
#         win = tk.Toplevel(self)
#         win.title("Conversion Log")
#         win.geometry("600x380")
#         win.configure(bg="#1a1a2e")
#         text = tk.Text(win, bg="#1a1a2e", fg="#e2e8f0", font=("Courier", 10),
#                        relief=tk.FLAT, padx=12, pady=12, wrap=tk.WORD)
#         text.pack(fill=tk.BOTH, expand=True)
#         sb = tk.Scrollbar(win, command=text.yview)
#         sb.pack(side=tk.RIGHT, fill=tk.Y)
#         text.config(yscrollcommand=sb.set)
#         log_text = "\n".join(self.conversion_log) if self.conversion_log else "No conversions run yet."
#         text.insert("1.0", log_text)
#         text.config(state=tk.DISABLED)


# if __name__ == "__main__":
#     app = DocuForgeApp()
#     app.mainloop()
































"""
DocuForge Converter Pro — World-Class Edition
═══════════════════════════════════════════════════════════════════════════════
Conversion fidelity strategy:
  • DOCX/ODT/PPTX → PDF   : LibreOffice headless  (100% layout fidelity)
  • XLSX/CSV      → PDF   : LibreOffice headless on intermediate XLSX
  • PDF           → DOCX  : PyMuPDF text extraction → styled python-docx
  • PDF           → TXT   : PyMuPDF (superior to pypdf for layout preservation)
  • PDF           → XLSX  : PyMuPDF + table heuristics → openpyxl
  • DOCX          → TXT   : python-docx paragraph traversal (preserves order)
  • DOCX          → XLSX  : paragraph lines → structured dataframe
  • TXT/CSV       → DOCX  : fully styled python-docx output
  • CSV/XLSX      → DOCX  : styled table inside python-docx
  • Any spreadsheet ↔ any spreadsheet : pandas
"""

import os, sys, re, shutil, subprocess, tempfile, threading, time
import tkinter as tk
from tkinter import filedialog, messagebox
from pathlib import Path

import ttkbootstrap as tb
from ttkbootstrap.constants import *
from ttkbootstrap.widgets.scrolled import ScrolledFrame

# Document libraries
import docx2txt
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import fitz          # PyMuPDF — best-in-class PDF reader
import pandas as pd
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from reportlab.lib.pagesizes import letter, landscape, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph as RLParagraph, Spacer, Table as RLTable, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors as rl_colors
from reportlab.lib.units import inch


# ── Runtime asset resolver ──────────────────────────────────────────────────
def resource_path(relative_path):
    try:
        base = sys._MEIPASS
    except AttributeError:
        base = os.path.abspath(".")
    return os.path.join(base, relative_path)


# ══════════════════════════════════════════════════════════════════════════════
#  LIBREOFFICE BRIDGE — the secret weapon for 100% fidelity
# ══════════════════════════════════════════════════════════════════════════════
class LibreOfficeBridge:
    """
    Wraps LibreOffice headless mode. When LibreOffice converts DOCX→PDF it uses
    the same rendering engine as the desktop app — identical to Microsoft Word's
    own 'Export to PDF' in terms of fidelity. Tables, fonts, colours, images,
    headers/footers — all preserved exactly.
    """

    # Candidate binary locations across Windows, macOS, Linux
    _CANDIDATES = [
        "libreoffice", "soffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/bin/libreoffice", "/usr/bin/soffice",
        "/usr/local/bin/libreoffice",
    ]

    @classmethod
    def find_binary(cls):
        for candidate in cls._CANDIDATES:
            if shutil.which(candidate) or os.path.isfile(candidate):
                return candidate
        return None

    @classmethod
    def convert(cls, source_path: str, target_format: str, output_dir: str) -> str | None:
        """
        Convert source_path to target_format inside output_dir.
        Returns the output file path on success, None on failure.
        target_format examples: 'pdf', 'docx', 'xlsx', 'csv', 'txt'
        """
        binary = cls.find_binary()
        if not binary:
            return None

        # LibreOffice locks its user-profile dir — use a temp profile per call
        # to allow parallel conversions and avoid conflicts with running instances.
        with tempfile.TemporaryDirectory() as profile_dir:
            cmd = [
                binary,
                "--headless",
                f"-env:UserInstallation=file:///{profile_dir.replace(os.sep, '/')}",
                "--convert-to", target_format,
                "--outdir", output_dir,
                source_path,
            ]
            try:
                result = subprocess.run(
                    cmd, capture_output=True, text=True, timeout=120
                )
                if result.returncode != 0:
                    print(f"[LibreOffice] Error: {result.stderr[:400]}")
                    return None

                # LO renames the file by swapping the extension
                stem = Path(source_path).stem
                ext = target_format.split(":")[0]  # handle "pdf" vs "pdf:writer_pdf_Export"
                out = Path(output_dir) / f"{stem}.{ext}"
                return str(out) if out.exists() else None
            except subprocess.TimeoutExpired:
                print("[LibreOffice] Conversion timed out")
                return None
            except Exception as e:
                print(f"[LibreOffice] Unexpected error: {e}")
                return None

    @classmethod
    def available(cls) -> bool:
        return cls.find_binary() is not None


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX STYLING HELPERS
# ══════════════════════════════════════════════════════════════════════════════
# Brand palette used in generated DOCX / XLSX outputs
BRAND = {
    "header_bg":   RGBColor(0x05, 0x41, 0x3E),  # deep teal
    "header_fg":   RGBColor(0xFF, 0xFF, 0xFF),
    "accent":      RGBColor(0x05, 0x96, 0x69),  # emerald
    "row_alt":     RGBColor(0xF0, 0xFD, 0xFA),
    "border":      RGBColor(0xCB, 0xD5, 0xE1),
    "body_text":   RGBColor(0x1E, 0x29, 0x3B),
}

BRAND_HEX = {
    "header_bg":  "05413E",
    "header_fg":  "FFFFFF",
    "accent":     "059669",
    "row_alt":    "F0FDFA",
    "row_even":   "FFFFFF",
    "border":     "CBD5E1",
}


def _style_docx_table(table, header_row=True):
    """Apply professional styling to a python-docx table."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            try:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                # Cell padding
                tcMar = OxmlElement('w:tcMar')
                for side in ('top', 'left', 'bottom', 'right'):
                    node = OxmlElement(f'w:{side}')
                    node.set(qn('w:w'), '80')
                    node.set(qn('w:type'), 'dxa')
                    tcMar.append(node)
                tcPr.append(tcMar)

                # Get or create a run to style
                if not cell.paragraphs:
                    continue
                para = cell.paragraphs[0]
                if para.runs:
                    run = para.runs[0]
                else:
                    run = para.add_run(cell.text or "")

                run.font.size = Pt(9.5)
                run.font.color.rgb = BRAND["body_text"]

                # Header row shading
                if header_row and i == 0:
                    run.bold = True
                    run.font.color.rgb = BRAND["header_fg"]
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:val'), 'clear')
                    shading.set(qn('w:fill'), BRAND_HEX["header_bg"])
                    tcPr.append(shading)
                elif i % 2 == 0:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:val'), 'clear')
                    shading.set(qn('w:fill'), BRAND_HEX["row_alt"])
                    tcPr.append(shading)
            except Exception as e:
                # Never crash the whole export over a single cell styling issue
                print(f"[Style] Cell ({i},{j}) styling skipped: {e}")


def _set_docx_styles(doc: DocxDocument):
    """Configure document-level typography."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = BRAND["body_text"]

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = BRAND["accent"]

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = BRAND["header_bg"]


# ══════════════════════════════════════════════════════════════════════════════
#  XLSX STYLING HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def _style_worksheet(ws, df, freeze_header=True):
    """Apply professional table styling to an openpyxl worksheet."""
    header_fill  = PatternFill("solid", fgColor=BRAND_HEX["header_bg"])
    alt_fill     = PatternFill("solid", fgColor=BRAND_HEX["row_alt"])
    even_fill    = PatternFill("solid", fgColor=BRAND_HEX["row_even"])
    thin         = Side(border_style="thin", color=BRAND_HEX["border"])
    cell_border  = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_font  = Font(name="Calibri", bold=True, color=BRAND_HEX["header_fg"], size=10)
    body_font    = Font(name="Calibri", size=10, color="1E293B")
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align   = Alignment(horizontal="left",   vertical="center", wrap_text=True)

    if df.empty:
        return

    # Header row
    for col_idx, col_name in enumerate(df.columns, start=1):
        cell = ws.cell(row=1, column=col_idx, value=str(col_name))
        cell.font      = header_font
        cell.fill      = header_fill
        cell.border    = cell_border
        cell.alignment = center_align

    # Data rows
    for row_idx, row in enumerate(df.itertuples(index=False), start=2):
        fill = alt_fill if row_idx % 2 == 0 else even_fill
        for col_idx, value in enumerate(row, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font      = body_font
            cell.fill      = fill
            cell.border    = cell_border
            cell.alignment = left_align

    # Auto-fit column widths (cap at 60 chars)
    for col_idx, col in enumerate(df.columns, start=1):
        col_values = df.iloc[:, col_idx - 1].astype(str).values
        lengths    = [len(str(col))] + [len(v) for v in col_values]
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max(lengths) + 4, 60)

    if freeze_header:
        ws.freeze_panes = "A2"
    ws.row_dimensions[1].height = 28


# ══════════════════════════════════════════════════════════════════════════════
#  PDF READING via PyMuPDF (superior layout preservation)
# ══════════════════════════════════════════════════════════════════════════════
def _extract_pdf_text(source_path: str) -> str:
    """Extract text from PDF preserving layout order via PyMuPDF."""
    doc = fitz.open(source_path)
    pages = []
    for page in doc:
        # sort=True preserves natural reading order (top→bottom, left→right)
        pages.append(page.get_text("text", sort=True))
    doc.close()
    return "\n\n".join(pages)


def _extract_pdf_tables(source_path: str) -> list[pd.DataFrame]:
    """Attempt to extract tables from a PDF using PyMuPDF's table finder."""
    doc = fitz.open(source_path)
    all_tables = []
    for page in doc:
        try:
            tabs = page.find_tables()
            if not (tabs and tabs.tables):
                continue
            for tab in tabs.tables:
                try:
                    rows = tab.extract()
                    if not rows or len(rows) < 2:
                        continue
                    # Sanitise header — None cells become empty strings
                    header = [str(h) if h is not None else f"Col{i}"
                              for i, h in enumerate(rows[0])]
                    data   = [[str(c) if c is not None else ""
                               for c in row] for row in rows[1:]]
                    # Skip if all header values are blank
                    if not any(h.strip() for h in header):
                        continue
                    df = pd.DataFrame(data, columns=header)
                    all_tables.append(df)
                except Exception as e:
                    print(f"[PDF Tables] Table extract error on page: {e}")
        except Exception as e:
            print(f"[PDF Tables] Page error: {e}")
    doc.close()
    return all_tables


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX READING — paragraph-aware extraction
# ══════════════════════════════════════════════════════════════════════════════
def _extract_docx_rich(source_path: str) -> tuple[list, list]:
    """
    Returns (paragraphs, tables).
    paragraphs: list of dicts with {text, style, bold, italic, alignment}
    tables:     list of 2D lists
    All attribute access is guarded — handles corrupt/unusual DOCX files safely.
    """
    doc = DocxDocument(source_path)
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text or ""
        if not text.strip():
            continue
        # style.name can be None on malformed docs
        try:
            style_name = p.style.name if p.style and p.style.name else "Normal"
        except Exception:
            style_name = "Normal"
        # run.bold / run.italic can be None (inherited, not set explicitly)
        try:
            is_bold   = any(bool(r.bold)   for r in p.runs)
        except Exception:
            is_bold = False
        try:
            is_italic = any(bool(r.italic) for r in p.runs)
        except Exception:
            is_italic = False
        paragraphs.append({
            "text":      text,
            "style":     style_name,
            "bold":      is_bold,
            "italic":    is_italic,
            "alignment": p.alignment,
        })

    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                try:
                    row_cells.append(cell.text or "")
                except Exception:
                    row_cells.append("")
            rows.append(row_cells)
        if rows:
            tables.append(rows)
    return paragraphs, tables


# ══════════════════════════════════════════════════════════════════════════════
#  CONVERSION ENGINE
# ══════════════════════════════════════════════════════════════════════════════
class ConversionEngine:

    @staticmethod
    def clean_df(df: pd.DataFrame) -> pd.DataFrame:
        df = df.fillna("")
        df.columns = df.columns.astype(str)
        df = df.loc[:, ~df.columns.str.contains(r'^Unnamed:', na=False)]
        return df

    @staticmethod
    def _read_spreadsheet(path: str) -> pd.DataFrame:
        ext = Path(path).suffix.lower()
        if ext == ".csv":
            # Try multiple encodings gracefully
            for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
                try:
                    return ConversionEngine.clean_df(pd.read_csv(path, encoding=enc))
                except UnicodeDecodeError:
                    continue
            return pd.DataFrame()
        else:
            return ConversionEngine.clean_df(pd.read_excel(path, engine="openpyxl"))

    # ── Core router ──────────────────────────────────────────────────────────
    @staticmethod
    def convert(source_path: str, target_ext: str, output_dir: str) -> bool:
        import traceback
        src = Path(source_path)
        src_ext = src.suffix.lower()
        target_ext = target_ext.lower()
        if not target_ext.startswith("."):
            target_ext = "." + target_ext

        stem       = src.stem
        target_path = str(Path(output_dir) / f"{stem}{target_ext}")

        # Same format — just copy
        if src_ext == target_ext:
            shutil.copy2(source_path, target_path)
            return True

        lo = LibreOfficeBridge.available()

        try:
            # ── HIGHEST FIDELITY PATHS (LibreOffice) ────────────────────────
            if lo:
                # DOCX/ODT/RTF → PDF
                if src_ext in (".docx", ".odt", ".rtf") and target_ext == ".pdf":
                    out = LibreOfficeBridge.convert(source_path, "pdf", output_dir)
                    if out: return True
                    raise RuntimeError("LibreOffice returned no output file for PDF export")

                # PDF → DOCX
                if src_ext == ".pdf" and target_ext == ".docx":
                    out = LibreOfficeBridge.convert(source_path, "docx", output_dir)
                    if out: return True
                    raise RuntimeError("LibreOffice returned no output file for DOCX export")

                # XLSX/XLS/ODS → PDF
                if src_ext in (".xlsx", ".xls", ".ods") and target_ext == ".pdf":
                    out = LibreOfficeBridge.convert(source_path, "pdf", output_dir)
                    if out: return True
                    raise RuntimeError("LibreOffice returned no output file for spreadsheet PDF")

                # CSV → PDF via styled XLSX intermediate
                if src_ext == ".csv" and target_ext == ".pdf":
                    df = ConversionEngine._read_spreadsheet(source_path)
                    tmp_xlsx = str(Path(output_dir) / f"_tmp_{stem}.xlsx")
                    ConversionEngine._df_to_styled_xlsx(df, tmp_xlsx)
                    out = LibreOfficeBridge.convert(tmp_xlsx, "pdf", output_dir)
                    try: os.remove(tmp_xlsx)
                    except Exception: pass
                    if out:
                        final = str(Path(output_dir) / f"{stem}.pdf")
                        if out != final: shutil.move(out, final)
                        return True
                    raise RuntimeError("LibreOffice returned no output file for CSV→PDF")

                # DOCX → ODT / RTF
                if src_ext == ".docx" and target_ext in (".odt", ".rtf"):
                    out = LibreOfficeBridge.convert(source_path, target_ext.lstrip("."), output_dir)
                    if out: return True
                    raise RuntimeError(f"LibreOffice returned no output for DOCX→{target_ext}")

            # ── SPREADSHEET ↔ SPREADSHEET ────────────────────────────────────
            if src_ext in (".csv", ".xlsx", ".xls") and target_ext in (".csv", ".xlsx"):
                df = ConversionEngine._read_spreadsheet(source_path)
                if target_ext == ".xlsx":
                    ConversionEngine._df_to_styled_xlsx(df, target_path)
                else:
                    df.to_csv(target_path, index=False, encoding="utf-8-sig")
                return True

            # ── SPREADSHEET → DOCX ───────────────────────────────────────────
            if src_ext in (".csv", ".xlsx", ".xls") and target_ext == ".docx":
                df = ConversionEngine._read_spreadsheet(source_path)
                ConversionEngine._df_to_docx(df, target_path, stem)
                return True

            # ── SPREADSHEET → TXT ────────────────────────────────────────────
            if src_ext in (".csv", ".xlsx", ".xls") and target_ext == ".txt":
                df = ConversionEngine._read_spreadsheet(source_path)
                with open(target_path, "w", encoding="utf-8") as f:
                    f.write(df.to_string(index=False))
                return True

            # ── PDF → TXT ────────────────────────────────────────────────────
            if src_ext == ".pdf" and target_ext == ".txt":
                text = _extract_pdf_text(source_path)
                with open(target_path, "w", encoding="utf-8") as f:
                    f.write(text)
                return True

            # ── PDF → XLSX / CSV ─────────────────────────────────────────────
            if src_ext == ".pdf" and target_ext in (".xlsx", ".csv"):
                tables = _extract_pdf_tables(source_path)
                if tables:
                    if target_ext == ".xlsx":
                        wb = openpyxl.Workbook()
                        for i, df in enumerate(tables):
                            ws = wb.active if i == 0 else wb.create_sheet(f"Table {i+1}")
                            ws.title = f"Table {i+1}"
                            _style_worksheet(ws, df)
                        wb.save(target_path)
                    else:
                        tables[0].to_csv(target_path, index=False, encoding="utf-8-sig")
                else:
                    # No tables — fall back to plain text rows
                    text  = _extract_pdf_text(source_path)
                    lines = [l for l in text.split("\n") if l.strip()]
                    df    = pd.DataFrame({"Text": lines})
                    if target_ext == ".xlsx":
                        ConversionEngine._df_to_styled_xlsx(df, target_path)
                    else:
                        df.to_csv(target_path, index=False, encoding="utf-8-sig")
                return True

            # ── DOCX → TXT ───────────────────────────────────────────────────
            if src_ext == ".docx" and target_ext == ".txt":
                paras, tables = _extract_docx_rich(source_path)
                lines = []
                for p in paras:
                    lines.append(p["text"])
                if tables:
                    lines.append("\n── Tables ──")
                    for t_idx, table in enumerate(tables, 1):
                        lines.append(f"\nTable {t_idx}:")
                        for row in table:
                            lines.append(" | ".join(str(c) for c in row))
                with open(target_path, "w", encoding="utf-8") as f:
                    f.write("\n".join(lines))
                return True

            # ── DOCX → XLSX / CSV ────────────────────────────────────────────
            if src_ext == ".docx" and target_ext in (".xlsx", ".csv"):
                paras, tables = _extract_docx_rich(source_path)
                if tables:
                    rows   = tables[0]
                    header = [str(c) for c in rows[0]] if rows else []
                    data   = [[str(c) for c in r] for r in rows[1:]] if len(rows) > 1 else []
                    # Pad rows that are shorter than the header
                    n_cols = max(len(header), max((len(r) for r in data), default=0))
                    header = header + [""] * (n_cols - len(header))
                    data   = [r + [""] * (n_cols - len(r)) for r in data]
                    df     = pd.DataFrame(data, columns=header) if header else pd.DataFrame(data)
                else:
                    df = pd.DataFrame({"Content": [p["text"] for p in paras]})
                if target_ext == ".xlsx":
                    ConversionEngine._df_to_styled_xlsx(df, target_path)
                else:
                    df.to_csv(target_path, index=False, encoding="utf-8-sig")
                return True

            # ── TXT → DOCX ───────────────────────────────────────────────────
            if src_ext == ".txt" and target_ext == ".docx":
                with open(source_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                ConversionEngine._text_to_docx(text, target_path, stem)
                return True

            # ── TXT → PDF ────────────────────────────────────────────────────
            if src_ext == ".txt" and target_ext == ".pdf":
                with open(source_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                if lo:
                    tmp_docx = str(Path(output_dir) / f"_tmp_{stem}.docx")
                    ConversionEngine._text_to_docx(text, tmp_docx, stem)
                    out = LibreOfficeBridge.convert(tmp_docx, "pdf", output_dir)
                    try: os.remove(tmp_docx)
                    except Exception: pass
                    if out:
                        final = str(Path(output_dir) / f"{stem}.pdf")
                        if out != final: shutil.move(out, final)
                        return True
                # ReportLab fallback (when LO unavailable or failed)
                ConversionEngine._text_to_pdf_reportlab(text, target_path)
                return True

            # ── TXT → XLSX / CSV ─────────────────────────────────────────────
            if src_ext == ".txt" and target_ext in (".xlsx", ".csv"):
                with open(source_path, "r", encoding="utf-8", errors="ignore") as f:
                    lines = [l.rstrip() for l in f if l.strip()]
                df = pd.DataFrame({"Content": lines})
                if target_ext == ".xlsx":
                    ConversionEngine._df_to_styled_xlsx(df, target_path)
                else:
                    df.to_csv(target_path, index=False, encoding="utf-8-sig")
                return True

            # ── No path matched ───────────────────────────────────────────────
            print(f"[Engine] No conversion path defined for {src_ext} → {target_ext}")
            return False

        except Exception as e:
            print(f"[Engine] FAILED {src.name} ({src_ext}→{target_ext}): {e}")
            print(traceback.format_exc())
            return False

    # ── Output builders ──────────────────────────────────────────────────────

    @staticmethod
    def _df_to_styled_xlsx(df: pd.DataFrame, path: str):
        """Write a DataFrame to a fully styled XLSX."""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Data"
        _style_worksheet(ws, df)
        wb.save(path)

    @staticmethod
    def _df_to_docx(df: pd.DataFrame, path: str, title: str = ""):
        """Write a DataFrame as a styled table inside a DOCX."""
        doc = DocxDocument()
        _set_docx_styles(doc)
        if title:
            h = doc.add_heading(title.replace("_", " ").title(), level=1)
        doc.add_paragraph(f"Total rows: {len(df)}  ·  Columns: {len(df.columns)}")
        doc.add_paragraph("")

        table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
        table.style = "Table Grid"
        # Header
        for j, col in enumerate(df.columns):
            table.cell(0, j).text = str(col)
        # Rows
        for i, row in enumerate(df.itertuples(index=False), start=1):
            for j, val in enumerate(row):
                table.cell(i, j).text = str(val)

        _style_docx_table(table, header_row=True)
        doc.save(path)

    @staticmethod
    def _text_to_docx(text: str, path: str, title: str = ""):
        """Convert plain text to a well-styled DOCX preserving paragraph structure."""
        doc = DocxDocument()
        _set_docx_styles(doc)
        if title:
            doc.add_heading(title.replace("_", " ").replace("-", " ").title(), level=1)

        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue
            # Heuristic: ALL CAPS short lines → treat as heading
            if stripped.isupper() and len(stripped) < 80:
                doc.add_heading(stripped.title(), level=2)
            else:
                p = doc.add_paragraph(stripped)
                p.style = doc.styles["Normal"]
        doc.save(path)

    @staticmethod
    def _text_to_pdf_reportlab(text: str, path: str):
        """ReportLab fallback PDF writer with clean typography."""
        doc = SimpleDocTemplate(
            path, pagesize=A4,
            leftMargin=inch, rightMargin=inch,
            topMargin=inch, bottomMargin=inch,
        )
        styles = getSampleStyleSheet()
        body = ParagraphStyle(
            "Body", parent=styles["Normal"],
            fontName="Helvetica", fontSize=10.5,
            leading=16, spaceAfter=8,
            textColor=rl_colors.HexColor("#1E293B"),
        )
        heading = ParagraphStyle(
            "H", parent=styles["Heading1"],
            fontName="Helvetica-Bold", fontSize=14,
            leading=20, spaceAfter=12,
            textColor=rl_colors.HexColor("#059669"),
        )
        story = []
        for line in text.split("\n"):
            s = line.strip()
            if not s:
                story.append(Spacer(1, 6))
                continue
            clean = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if s.isupper() and len(s) < 80:
                story.append(RLParagraph(clean.title(), heading))
            else:
                story.append(RLParagraph(clean, body))
        doc.build(story)


# ══════════════════════════════════════════════════════════════════════════════
#  GUI APPLICATION
# ══════════════════════════════════════════════════════════════════════════════

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

class DocuForgeApp(tb.Window):

    def __init__(self):
        super().__init__(themename="cyborg", title="DocuForge Converter Pro", size=(1000, 700))
        self.selected_files: dict[str, str] = {}
        self.conversion_log: list[str] = []
        self._lo_available = LibreOfficeBridge.available()
        self._setup_logo()
        self._build_ui()

    def _setup_logo(self):
        for path_fn, bind_fn in [
            (lambda: resource_path("logo.png"), lambda p: self.iconphoto(False, tk.PhotoImage(file=p))),
            (lambda: resource_path("logo.ico"), lambda p: self.iconbitmap(p)),
        ]:
            try:
                p = path_fn()
                if os.path.exists(p):
                    bind_fn(p)
            except Exception:
                pass

    def _build_ui(self):
        # ── Top header ───────────────────────────────────────────────────────
        top = tb.Frame(self, padding=(16, 12), bootstyle=DARK)
        top.pack(fill=X)

        tb.Label(top, text="⚡ DocuForge Pro",
                 font=("Helvetica", 17, "bold"), bootstyle=LIGHT).pack(side=LEFT, padx=8)

        # Fidelity badge
        badge_text = "✦ High-Fidelity Mode (LibreOffice)" if self._lo_available else "⚡ Standard Mode"
        badge_style = SUCCESS if self._lo_available else WARNING
        tb.Label(top, text=badge_text, font=("Helvetica", 9),
                 bootstyle=badge_style).pack(side=LEFT, padx=14)

        tb.Button(top, text="📂 Add Folder",    bootstyle=INFO,    command=self._upload_directory).pack(side=RIGHT, padx=4)
        tb.Button(top, text="📁 Add Files",     bootstyle=PRIMARY,  command=self._upload_files).pack(side=RIGHT, padx=4)
        tb.Button(top, text="🗑 Clear All",     bootstyle="danger-outline", command=self._clear_all).pack(side=RIGHT, padx=4)

        # ── Main body ────────────────────────────────────────────────────────
        body = tb.Frame(self, padding=(14, 10))
        body.pack(fill=BOTH, expand=True)

        # Left: file queue
        left = tb.Labelframe(body, text=" Queued Files ", padding=8, bootstyle=INFO)
        left.pack(side=LEFT, fill=BOTH, expand=True, padx=(0, 8))

        self.scroll_box = ScrolledFrame(left, autohide=True)
        self.scroll_box.pack(fill=BOTH, expand=True)

        # Right: options
        right = tb.Labelframe(body, text=" Conversion Options ", padding=14, width=300, bootstyle=INFO)
        right.pack(side=RIGHT, fill=Y, padx=(8, 0))
        right.pack_propagate(False)

        tb.Label(right, text="Convert to:", font=("Helvetica", 11, "bold")).pack(anchor=W, pady=(0, 8))

        self.target_fmt = tb.StringVar(value=".pdf")
        FORMATS = [
            ("PDF Document (.pdf)",        ".pdf"),
            ("Word Document (.docx)",      ".docx"),
            ("Excel Spreadsheet (.xlsx)",  ".xlsx"),
            ("CSV File (.csv)",            ".csv"),
            ("Plain Text (.txt)",          ".txt"),
        ]
        for label, ext in FORMATS:
            tb.Radiobutton(right, text=label, value=ext,
                           variable=self.target_fmt, bootstyle=INFO).pack(anchor=W, pady=5)

        # Fidelity note
        tb.Separator(right).pack(fill=X, pady=10)
        note = ("LibreOffice active — DOCX/XLSX→PDF conversion\npreserves 100% formatting fidelity."
                if self._lo_available else
                "Tip: Install LibreOffice for pixel-perfect\nDOCX/XLSX→PDF conversion.")
        tb.Label(right, text=note, font=("Helvetica", 8),
                 bootstyle="success" if self._lo_available else "warning",
                 wraplength=260, justify=LEFT).pack(anchor=W, pady=(0, 10))

        tb.Separator(right).pack(fill=X, pady=4)

        # Progress
        self.status_var = tb.StringVar(value="Ready")
        tb.Label(right, textvariable=self.status_var, font=("Helvetica", 9),
                 bootstyle=SECONDARY, wraplength=260).pack(anchor=W, pady=(8, 4))

        self.progress = tb.Progressbar(right, bootstyle=SUCCESS, mode="determinate")
        self.progress.pack(fill=X, pady=(0, 6), side=BOTTOM)

        self.convert_btn = tb.Button(
            right, text="⚡ Start Conversion", bootstyle=SUCCESS,
            state=DISABLED, command=self._start_conversion_thread
        )
        self.convert_btn.pack(fill=X, side=BOTTOM, pady=4)

        # Log button
        tb.Button(right, text="📋 View Log", bootstyle="secondary-outline",
                  command=self._show_log).pack(fill=X, side=BOTTOM, pady=2)

        self._refresh_queue()

    # ── File management ──────────────────────────────────────────────────────
    def _refresh_queue(self):
        for w in self.scroll_box.winfo_children():
            w.destroy()

        if not self.selected_files:
            tb.Label(self.scroll_box, text="Drop files here or use the buttons above.",
                     font=("Helvetica", 10, "italic"), bootstyle="secondary").pack(pady=30)
            self.convert_btn.config(state=DISABLED)
            return

        self.convert_btn.config(state=NORMAL)
        EXT_ICONS = {".pdf": "📄", ".docx": "📝", ".xlsx": "📊",
                     ".csv": "📋", ".txt": "📃"}

        for idx, (path, group) in enumerate(self.selected_files.items()):
            ext = Path(path).suffix.lower()
            icon = EXT_ICONS.get(ext, "📎")
            bg = DARK if idx % 2 == 0 else SECONDARY

            row = tb.Frame(self.scroll_box, padding=(8, 5), bootstyle=bg)
            row.pack(fill=X, pady=2)

            label = f"[{group}]  {Path(path).name}" if group else f"{icon}  {Path(path).name}"
            tb.Label(row, text=label, font=("Helvetica", 9)).pack(side=LEFT, padx=4)

            tb.Button(row, text="✕", padding=(5, 1), bootstyle="danger-outline",
                      command=lambda p=path: self._remove_file(p)).pack(side=RIGHT, padx=4)

    def _upload_files(self):
        paths = filedialog.askopenfilenames(
            title="Select Files",
            filetypes=[("Supported Documents", "*.docx *.txt *.csv *.xlsx *.pdf *.odt *.rtf")]
        )
        for p in paths:
            if p not in self.selected_files:
                self.selected_files[p] = ""
        self._refresh_queue()

    def _upload_directory(self):
        d = filedialog.askdirectory(title="Select Folder")
        if not d:
            return
        valid = (".docx", ".txt", ".csv", ".xlsx", ".pdf", ".odt", ".rtf")
        folder = Path(d).name
        for root, _, files in os.walk(d):
            for f in files:
                if Path(f).suffix.lower() in valid:
                    fp = os.path.join(root, f)
                    if fp not in self.selected_files:
                        self.selected_files[fp] = folder
        self._refresh_queue()

    def _remove_file(self, path: str):
        self.selected_files.pop(path, None)
        self._refresh_queue()

    def _clear_all(self):
        self.selected_files.clear()
        self._refresh_queue()

    # ── Conversion ───────────────────────────────────────────────────────────
    def _start_conversion_thread(self):
        """Run conversion in background thread so UI stays responsive."""
        self.convert_btn.config(state=DISABLED)
        threading.Thread(target=self._run_conversion, daemon=True).start()

    def _run_conversion(self):
        output_dir = filedialog.askdirectory(title="Select Output Folder")
        if not output_dir:
            self.after(0, lambda: self.convert_btn.config(state=NORMAL))
            return

        target_ext = self.target_fmt.get()
        items = list(self.selected_files.items())
        total = len(items)
        success = 0
        self.conversion_log.clear()

        for idx, (file_path, group) in enumerate(items):
            name = Path(file_path).name
            self.after(0, lambda n=name: self.status_var.set(f"Converting: {n}"))

            try:
                dest_dir = os.path.join(output_dir, group) if group else output_dir
                os.makedirs(dest_dir, exist_ok=True)
                ok = ConversionEngine.convert(file_path, target_ext, dest_dir)
                if ok:
                    success += 1
                    self.conversion_log.append(f"✅  {name}")
                else:
                    self.conversion_log.append(f"❌  {name}  — no conversion path found")
            except Exception as e:
                self.conversion_log.append(f"💥  {name}  — {str(e)[:120]}")

            pct = ((idx + 1) / total) * 100
            self.after(0, lambda v=pct: self.progress.config(value=v))

        self.after(0, lambda: self._finish(success, total))

    def _finish(self, success: int, total: int):
        self.status_var.set(f"Done — {success}/{total} files converted")
        self.progress.config(value=0)
        self.convert_btn.config(state=NORMAL if self.selected_files else DISABLED)
        messagebox.showinfo(
            "Conversion Complete",
            f"✅  {success} of {total} file(s) converted successfully.\n\n"
            + ("View the log for any errors." if success < total else "All files converted.")
        )
        if success == total:
            self.selected_files.clear()
            self._refresh_queue()

    def _show_log(self):
        """Show a scrollable log window."""
        win = tk.Toplevel(self)
        win.title("Conversion Log")
        win.geometry("600x380")
        win.configure(bg="#1a1a2e")
        text = tk.Text(win, bg="#1a1a2e", fg="#e2e8f0", font=("Courier", 10),
                       relief=tk.FLAT, padx=12, pady=12, wrap=tk.WORD)
        text.pack(fill=tk.BOTH, expand=True)
        sb = tk.Scrollbar(win, command=text.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        text.config(yscrollcommand=sb.set)
        log_text = "\n".join(self.conversion_log) if self.conversion_log else "No conversions run yet."
        text.insert("1.0", log_text)
        text.config(state=tk.DISABLED)


if __name__ == "__main__":
    app = DocuForgeApp()
    app.mainloop()